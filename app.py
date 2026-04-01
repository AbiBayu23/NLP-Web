import streamlit as st
import re
import math
import PyPDF2
from docx import Document
from deep_translator import GoogleTranslator
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
import pandas as pd
from collections import Counter
import altair as alt
import io
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from nltk.sentiment import SentimentIntensityAnalyzer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from langdetect import detect
import torch
import tempfile
import time
from pydub import AudioSegment
import streamlit as st
import pandas as pd
from bs4 import BeautifulSoup
from collections import Counter
import html
import numpy as np
import nltk
import torch
import functools

torch.load = functools.partial(torch.load, weights_only=False)

from ai_engine import (
    load_ai_model, 
    load_whisper_model, 
    dapatkan_sinonim, 
    SPACY_MODELS
)
def inisialisasi_sistem_ai():
    spacy_models = ["en_core_web_lg", "es_core_news_lg", "fr_core_news_lg", "de_core_news_lg"]
    
    
    if 'local_files' not in st.session_state:
        st.session_state.local_files = {}
    
    if 'intro_seen' not in st.session_state:
        st.balloons()
        st.markdown("""
            <div style='background: linear-gradient(135deg, #0EA5E9 0%, #2563EB 100%); padding: 40px; border-radius: 15px; color: white; margin-bottom: 30px; text-align: center; box-shadow: 0 10px 25px -5px rgba(37, 99, 235, 0.4);'>
                <h1 style='color: white; font-size: 3rem; margin-bottom: 10px;'>Selamat Datang di AI-Explorer ☁️</h1>
                <p style='font-size: 1.2rem; opacity: 0.9;'>Platform Analisis Linguistik & Korpus Tercanggih untuk Dokumen & Suara</p>
                <hr style='border: 1px solid rgba(255,255,255,0.2); margin: 25px 0;'>
                <div style='display: flex; justify-content: center; gap: 20px; flex-wrap: wrap;'>
                    <div style='background: rgba(255,255,255,0.1); padding: 10px; border-radius: 8px;'>🔍 Lemmatization</div>
                    <div style='background: rgba(255,255,255,0.1); padding: 10px; border-radius: 8px;'>🎙️ Auto-Transcription</div>
                    <div style='background: rgba(255,255,255,0.1); padding: 10px; border-radius: 8px;'>📊 Visual Analytics</div>
                    <div style='background: rgba(255,255,255,0.1); padding: 10px; border-radius: 8px;'>📖 Interlinear Gloss</div>
                </div>
            </div>
        """, unsafe_allow_html=True)

        progress_bar = st.progress(0)
        status_text = st.empty()

        res_list = ['punkt', 'wordnet', 'omw-1.4', 'vader_lexicon', 'averaged_perceptron_tagger', 'punkt_tab']
        for i, res in enumerate(res_list):
            nltk.download(res, quiet=True)
            progress_bar.progress(10 + int((i+1)/len(res_list) * 20))

        import spacy
        for i, model in enumerate(spacy_models):
            if not spacy.util.is_package(model):
                status_text.warning(f"📥 Mengunduh model: {model}...")
                spacy.cli.download(model)
            progress_bar.progress(30 + int((i+1)/len(spacy_models) * 40))

        for lang in ['id', 'en']:
            load_ai_model(lang)("Pemanasan.")
        load_whisper_model()
        
        progress_bar.progress(100)
        time.sleep(1)
        progress_bar.empty()
        status_text.empty()
        if st.button("Masuk ke Dashboard Analisis →", type="primary"):
            st.session_state.intro_seen = True
            st.rerun()
        st.stop()

if 'system_ready' not in st.session_state:
    inisialisasi_sistem_ai()
    st.session_state.system_ready = True

_original_load = torch.load
def _patched_load(*args, **kwargs):
    kwargs['weights_only'] = False
    return _original_load(*args, **kwargs)
torch.load = _patched_load

def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split('([0-9]+)', s)]

def clean_duplicated_start(source, target):
    s = source.strip()
    t = target.strip()
    if s.lower() == t.lower():
        return t   
    words = s.split()
    if not words:
        return t
    pattern_str = r'\s+'.join(re.escape(w) for w in words)
    pattern = re.compile(r'^\s*' + pattern_str + r'\s*', re.IGNORECASE)
    
    if pattern.search(t):
        cleaned = pattern.sub('', t, count=1).strip()
        cleaned = re.sub(r'^[:,\-]\s*', '', cleaned)
        if cleaned:
            return cleaned   
    return t

def process_eaf_ultra_clean(eaf_content):
    """
    Fungsi ini menerima teks string XML (hasil decode dari Streamlit) 
    dan mengembalikan Pandas DataFrame.
    """
    soup = BeautifulSoup(eaf_content, 'xml')

    parent_dict = {}
    for align in soup.find_all('ALIGNABLE_ANNOTATION'):
        p_id = align.get('ANNOTATION_ID')
        val = align.find('ANNOTATION_VALUE')
        if p_id and val:
            parent_dict[p_id] = val.text.strip()

    tier_data = {}
    for tier in soup.find_all('TIER'):
        t_id = tier.get('TIER_ID', '')
        tier_data[t_id] = {}
        for ref in tier.find_all('REF_ANNOTATION'):
            ref_id = ref.get('ANNOTATION_REF')
            val = ref.find('ANNOTATION_VALUE')
            if ref_id and val:
                word = val.text.strip()
                if ref_id in tier_data[t_id]:
                    tier_data[t_id][ref_id] = (tier_data[t_id][ref_id] + " " + word).replace("  ", " ")
                else:
                    tier_data[t_id][ref_id] = word

    rows = []
    for p_id, full_source in parent_dict.items():
        raw_targets = []
        catatan_texts = []
        for t_id, ref_dict in tier_data.items():
            if p_id not in ref_dict: continue
            text_val = ref_dict[p_id].strip()
            if "-note" in t_id.lower():
                catatan_texts.append(text_val)
            else:
                raw_targets.append(text_val)
                
        norm_source = re.sub(r'\s+', ' ', full_source.strip().lower())
        has_real_translation = any(re.sub(r'\s+', ' ', t.strip().lower()) != norm_source for t in raw_targets)

        filtered_targets = []
        for t in raw_targets:
            t_norm = re.sub(r'\s+', ' ', t.strip().lower())
            if t_norm == norm_source and has_real_translation:
                continue
            cleaned_t = clean_duplicated_start(full_source, t)
            if cleaned_t and cleaned_t not in filtered_targets:
                filtered_targets.append(cleaned_t)

        target_text = " ".join(filtered_targets).strip()
        target_text = clean_duplicated_start(full_source, target_text)
        if target_text and not target_text.endswith(('.', '?', '!')):
            target_text += '.'
            
        catatan_text = " | ".join(catatan_texts).strip()

        if target_text or catatan_text:
            rows.append({
                "ID_Unit": p_id,
                "Source_Sentence": full_source,
                "Target_Sentence": target_text,
                "Catatan": catatan_text
            })

    rows.sort(key=lambda x: natural_sort_key(x['ID_Unit']))

    if rows:
        return pd.DataFrame(rows)
    return None

# ==========================================
# 1. KONFIGURASI HALAMAN WEB & TEMA
# ==========================================
st.set_page_config(
    page_title="AI NLP DOCUMENT ANALYSIS", 
    page_icon="☁️", 
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
    <style>
        .stApp { background-color: #F8FAFC; }
        
        [data-testid="stFileUploadDropzone"] {
            background-color: #FFFFFF !important;
            border: 2px dashed #BAE6FD !important;
            border-radius: 10px !important;
        }
        [data-testid="stFileUploadDropzone"]:hover {
            border: 2px dashed #0EA5E9 !important;
            background-color: #F0F9FF !important;
        }        
        div[data-testid="stTextInput"] input { 
            color: #0F172A !important; 
            font-weight:500 !important; 
            background-color: #FFFFFF !important;
            border: 1px solid #CBD5E1 !important;
        }
        div[data-testid="stTextInput"] input:focus {
            border-color: #0EA5E9 !important;
            box-shadow: 0 0 0 1px #0EA5E9 !important;
        }        
        div[data-baseweb="select"] > div {
            background-color: #FFFFFF !important;
            border: 1px solid #CBD5E1 !important;
            cursor: pointer !important;
        }            
        div[data-testid="stVerticalBlockBorderWrapper"] div[data-testid="stSelectbox"] {
            max-width: 140px !important;
            margin-left: auto !important;
        }        
        div[data-baseweb="select"] input {
            caret-color: transparent !important;
            cursor: pointer !important;
        }        
        button[data-baseweb="tab"] {
            background-color: transparent !important;
        }
    </style>
    <div style='background-color:#E0F2FE; padding:20px; border-radius:12px; border-left: 8px solid #0EA5E9; margin-bottom:25px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);'>
        <h2 style='color:#0369A1; margin:0; font-weight: 800;'>☁️ NLP Analisis Dokumen</h2>
        <p style='color:#475569; margin:8px 0 0 0; font-size: 15px;'>Multi-File Management • Vsiualisasi Data • POS Search • Rangkuman Dokumen</p>
    </div>
""", unsafe_allow_html=True)

if 'notif_msg' in st.session_state and 'notif_time' in st.session_state:
    if time.time() - st.session_state.notif_time < 10:
        st.success(st.session_state.notif_msg, icon="✅")
    else:
        del st.session_state.notif_msg
        del st.session_state.notif_time

# ==========================================
# 2. KAMUS & DAFTAR REFERENSI
# ==========================================
DAFTAR_BAHASA = {
    'Indonesian': 'id', 'English': 'en', 'Spanish': 'es', 
    'French': 'fr', 'German': 'de'
}

MAP_SEMUA_POS = {
    "NOUN (Kata Benda)": "NOUN", "VERB (Kata Kerja)": "VERB", "ADJ (Kata Sifat)": "ADJ",
    "ADV (Kata Keterangan)": "ADV", "PROPN (Nama/Entitas)": "PROPN", "PRON (Kata Ganti)": "PRON",
    "ADP (Preposisi/Kata Depan)": "ADP", "DET (Penentu/Determiner)": "DET", "AUX (Kata Bantu)": "AUX",
    "NUM (Angka)": "NUM", "PART (Partikel)": "PART", "SCONJ (Konjungsi Subordinatif)": "SCONJ",
    "CCONJ (Konjungsi Koordinatif)": "CCONJ", "INTJ (Interjeksi/Seruan)": "INTJ"
}

deskripsi_pos = {
    "NOUN": "Merujuk pada benda, manusia, tempat, atau ide abstrak.",
    "VERB": "Menyatakan tindakan, proses, atau keadaan.",
    "ADJ": "Menjelaskan ciri, sifat, atau keadaan dari kata benda.",
    "ADV": "Memberikan keterangan tambahan pada verba atau adjektiva.",
    "PROPN": "Nama diri yang spesifik seperti nama orang, tempat, atau merek.",
    "PRON": "Kata yang menggantikan penyebutan benda atau orang.",
    "ADP": "Kata depan yang menunjukkan hubungan ruang atau waktu.",
    "DET": "Kata yang memperjelas atau membatasi kata benda.",
    "AUX": "Kata bantu yang mendampingi kata kerja utama.",
    "NUM": "Menunjukkan jumlah, kuantitas, atau urutan angka.",
    "PART": "Kata tugas yang memiliki fungsi gramatikal khusus.",
    "SCONJ": "Penghubung antara anak kalimat dan induk kalimat.",
    "CCONJ": "Penghubung dua unsur kalimat yang setara atau sejajar.",
    "INTJ": "Kata seru untuk mengungkapkan emosi, perasaan, atau sapaan."
}

Warna_POS_Utama = {
    'NOUN': '#2563EB', 'VERB': '#D97706', 'ADJ': '#059669', 'ADV': '#DC2626', 
    'PRON': '#7C3AED', 'PROPN': '#E11D48', 'ADP': '#475569', 'DET': '#0891B2', 
    'AUX': '#4F46E5', 'NUM': '#DB2777', 'PART': '#10B981', 'SCONJ': '#9333EA', 
    'CCONJ': '#C026D3', 'INTJ': '#EA580C'
}

Warna_NER = {
    'PERSON': '#F43F5E', 'ORG': '#8B5CF6', 'GPE': '#10B981', 'LOC': '#059669', 
    'DATE': '#F59E0B', 'TIME': '#FBBC05', 'MONEY': '#22C55E', 'PRODUCT': '#3B82F6'
}

# ==========================================
# 3. FUNGSI-FUNGSI HELPER & CACHING UI
# ==========================================

@st.cache_resource(show_spinner=False)
def get_cached_spacy_doc(text, lang):
    """Menyimpan hasil parsing SpaCy ke dalam cache memori agar pencarian berulang jadi instan."""
    nlp_model = load_ai_model(lang)
    return nlp_model(text)

def buat_indeks_dokumen(fname, df_or_text, nlp_model, file_type='txt'):
    """Membangun tabel metadata agar pencarian tidak perlu membaca ulang seluruh teks"""
    data_indeks = []
    if file_type == 'eaf':
        for _, row in df_or_text.iterrows():
            teks_sumber = str(row['Source_Sentence']).strip()
            if not teks_sumber: continue
            
            doc = nlp_model(teks_sumber)
            data_indeks.append({
                'ID_Segmen': row['ID_Unit'], 
                'File': fname,
                'Teks_Asli': teks_sumber,
                'Tokens': [t.text.lower() for t in doc],
                'Lemmas': [t.lemma_.lower() for t in doc],
                'POS_Tags': [t.pos_ for t in doc]
            })
    else:
        kalimat_list = nltk.sent_tokenize(df_or_text)
        for i, teks_kalimat in enumerate(kalimat_list):
            doc = nlp_model(teks_kalimat)
            data_indeks.append({
                'ID_Segmen': f"Segmen_{i+1}",
                'File': fname,
                'Teks_Asli': teks_kalimat.strip(),
                'Tokens': [t.text.lower() for t in doc],
                'Lemmas': [t.lemma_.lower() for t in doc],
                'POS_Tags': [t.pos_ for t in doc]
            }) 
    return pd.DataFrame(data_indeks)

@st.cache_data(show_spinner=False)
def dapatkan_data_visual(teks_terbatas, _nlp_model):
    doc_vis = _nlp_model(teks_terbatas)
    
    pos_counts = Counter([token.pos_ for token in doc_vis if token.pos_ not in ['SPACE', 'PUNCT', 'SYM', 'X']])
    df_pos = pd.DataFrame(pos_counts.most_common(), columns=['POS Tag', 'Jumlah Kata']) if pos_counts else pd.DataFrame()
    
    raw_words = re.findall(r'\b[a-z]+(?:-[a-z]+)*\b', teks_terbatas.lower())
    stop_words_spacy = _nlp_model.Defaults.stop_words
    daftar_hitam_kustom = {
        'https', 'http', 'doi', 'et', 'al', 'www', 'com', 'org', 
        'pdf', 'fig', 'figure', 'table', 'vol', 'pp', 'ieee', 
        'the', 'and', 'for', 'that', 'using', 'based'
    }
    
    words_all = [w for w in raw_words if len(w) > 2 and w not in stop_words_spacy and w not in daftar_hitam_kustom]

    word_counts = Counter(words_all).most_common()
    df_words = pd.DataFrame(word_counts, columns=['Kata', 'Frekuensi']) if word_counts else pd.DataFrame()
    df_cloud = " ".join(words_all)
    
    return df_pos, df_words, df_cloud

@st.cache_data(show_spinner=False)
def get_cached_wordcloud(text_data):
    wc = WordCloud(width=600, height=280, background_color='white', colormap='viridis', max_words=100).generate(text_data)
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.imshow(wc, interpolation='bilinear')
    ax.axis("off")
    return fig

@st.cache_resource(show_spinner=False)
def load_bart_summarizer():
    from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
    import torch
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model_name = "sshleifer/distilbart-cnn-12-6"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name).to(device)
    
    return tokenizer, model, device

def render_dependency_tree(text, nlp_model):
    import spacy
    doc = nlp_model(text)
    options = {
        "compact": False, "distance": 100, "arrow_stroke": 2, "arrow_width": 6,
        "color": "#0284C7", "bg": "#FFFFFF", "font": "Source Sans Pro"
    }
    svg_html = spacy.displacy.render(doc, style="dep", options=options, page=False)
    
    legenda = {
        "nsubj": "<b>Nominal Subject</b>: Pelaku/Subjek utama dalam kalimat.",
        "ROOT": "<b>Root</b>: Inti atau kata kerja utama dalam kalimat.",
        "obj / dobj": "<b>Object</b>: Penerima tindakan dari kata kerja.",
        "amod": "<b>Adjectival Modifier</b>: Kata sifat yang menjelaskan kata benda.",
        "advmod": "<b>Adverbial Modifier</b>: Keterangan tambahan (waktu/cara/tempat).",
        "det": "<b>Determiner</b>: Kata penentu seperti 'a', 'the', atau 'an'.",
        "prep": "<b>Preposition</b>: Kata depan yang menghubungkan frasa."
    }
    legenda_html = "".join([f"<li>{v}</li>" for v in legenda.values()])
    
    full_html = f"""
    <div style='border: 1px solid #BAE6FD; border-radius: 8px; background: #F8FAFC; padding: 15px; margin-top: 10px;'>
        <div style='overflow-x: auto; margin-bottom: 15px; background: white; border-radius: 6px; padding: 10px;'>
            {svg_html}
        </div>
        <div style='background: #E0F2FE; padding: 12px; border-radius: 6px;'>
            <p style='margin: 0 0 8px 0; font-weight: bold; color: #0369A1; font-size: 14px;'>💡 Cara Membaca Struktur (Dependency Tags):</p>
            <ul style='margin: 0; padding-left: 20px; font-size: 13px; color: #334155; line-height: 1.5;'>
                {legenda_html}
            </ul>
        </div>
    </div>
    """
    return full_html

def buat_word_sketch(kata_target, teks_mentah, nlp_model):
    doc = nlp_model(teks_mentah[:20000]) 
    target_lemma = kata_target.lower()
    
    modifiers = []
    verbs_subject = []
    verbs_object = []
    
    for token in doc:
        if token.lemma_.lower() == target_lemma and not token.is_stop:
            for child in token.children:
                if child.dep_ in ['amod', 'advmod', 'compound'] and len(child.text) > 2:
                    modifiers.append(child.lemma_.lower())
            
            if token.dep_ in ['nsubj', 'nsubjpass']:
                if token.head.pos_ == 'VERB':
                    verbs_subject.append(token.head.lemma_.lower())
                    
            if token.dep_ in ['dobj', 'pobj']:
                head_token = token.head if token.dep_ == 'dobj' else token.head.head
                if head_token.pos_ == 'VERB':
                    verbs_object.append(head_token.lemma_.lower())
                    
    return {
        "✨ Modifiers (Sifat/Penjelas)": Counter(modifiers).most_common(5),
        "🏃‍♂️ Sebagai Subjek (Melakukan)": Counter(verbs_subject).most_common(5),
        "🎯 Sebagai Objek (Dikenai)": Counter(verbs_object).most_common(5)
    }

def hitung_collocation(kata_target, teks_mentah, nlp_model, window=3):
    words = re.findall(r'\b[a-z]+(?:-[a-z]+)*\b', teks_mentah.lower())
    stop_words = nlp_model.Defaults.stop_words
    pasangan = []

    for i, w in enumerate(words):
        if w == kata_target.lower():
            start = max(0, i - window)
            end = min(len(words), i + window + 1)
            for j in range(start, end):
                if i != j: 
                    kandidat = words[j]
                    if len(kandidat) > 2 and kandidat not in stop_words:
                        pasangan.append(kandidat)
                        
    return Counter(pasangan).most_common(5)

def get_colored_pos_text(nlp_model, text):
    doc = nlp_model(text)
    html = "<div style='line-height: 2.5; padding: 15px; background-color: #FFFFFF; border-radius: 8px; border: 1px solid #E2E8F0; box-shadow: inset 0 2px 4px 0 rgba(0, 0, 0, 0.03);'>"
    
    for token in doc:
        if token.pos_ == 'SPACE':
            html += token.text.replace('\n', '<br>')
            continue
            
        if token.pos_ == 'PUNCT':
            html += f"<span style='font-size: 15px; margin-right: 4px;'>{token.text}</span>"
            continue

        bg_color = Warna_POS_Utama.get(token.pos_, '#94A3B8')
        html += f"<span style='background-color: {bg_color}; color: white; padding: 4px 10px; border-radius: 6px; margin: 3px; font-size: 14.5px; display: inline-block; font-weight:600; box-shadow: 0 1px 2px rgba(0,0,0,0.1);'>"
        html += f"{token.text} <span style='font-size: 10.5px; opacity: 0.9; margin-left: 4px; text-transform: uppercase;'>{token.pos_}</span></span> "
        
    html += "</div>"
    return html

def get_colored_ner_inline(nlp_model, text, target_label=None):
    doc = nlp_model(text)
    if not doc.ents: return text
    html = ""
    last_idx = 0
    for ent in doc.ents:
        html += text[last_idx:ent.start_char]
        bg_color = Warna_NER.get(ent.label_, "#3A4B62")
        opacity = "1.0" if (target_label is None or ent.label_ == target_label) else "0.4"
        html += f"<mark style='background-color: {bg_color}; color: white; padding: 2px 6px; border-radius: 4px; font-weight:600; font-size: 14.5px; opacity: {opacity}; box-shadow: 0 1px 2px rgba(0,0,0,0.1);'>{ent.text} <span style='font-size:10px; text-transform:uppercase;'>{ent.label_}</span></mark>"
        last_idx = ent.end_char
    html += text[last_idx:]
    return html

def extract_text(uploaded_file):
    text = ""
    file_extension = uploaded_file.name.split('.')[-1].lower()
    try:
        if file_extension == 'pdf':
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages: text += page.extract_text() + "\n"
        elif file_extension == 'docx':
            doc = Document(uploaded_file)
            for para in doc.paragraphs: text += para.text + "\n"
        elif file_extension == 'txt':
            text = uploaded_file.getvalue().decode("utf-8")
    except Exception as e:
        st.error(f"Gagal membaca file: {e}")
    return text

def transkripsi_suara_whisper(audio_file):
    model = load_whisper_model()
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{audio_file.name.split('.')[-1]}") as tmp:
        tmp.write(audio_file.getvalue())
        tmp_path = tmp.name
    
    result = model.transcribe(tmp_path)
    return result["text"], result["segments"]

def terjemahkan_teks_panjang(teks, target_lang_code):
    paragraf = teks.split('\n')
    hasil = []
    for p in paragraf:
        if p.strip():
            if len(p) > 4900:
                potongan = [p[i:i+4900] for i in range(0, len(p), 4900)]
                for pot in potongan:
                    hasil.append(GoogleTranslator(source='auto', target=target_lang_code).translate(pot))
            else:
                hasil.append(GoogleTranslator(source='auto', target=target_lang_code).translate(p))
        else:
            hasil.append("")
    return '\n'.join(hasil)

def bersihkan_teks_untuk_analisis(teks_dokumen):
    pola_referensi = re.compile(r'\n\s*(DAFTAR PUSTAKA|REFERENCES|BIBLIOGRAPHY|REFERENSI).*', re.IGNORECASE | re.DOTALL)
    teks = re.sub(pola_referensi, '', teks_dokumen)
    pola_awal = re.compile(r'.*?(?=BAB\s?I|INTRODUCTION|PENDAHULUAN)', re.IGNORECASE | re.DOTALL)
    if re.search(r'BAB\s?I|INTRODUCTION|PENDAHULUAN', teks, re.IGNORECASE):
        teks = re.sub(pola_awal, '', teks, count=1)
    teks = re.sub(r'(?m)^(Figure|Gambar|Table|Tabel|DOI|ISSN|Source).*$', '', teks)
    teks = re.sub(r'\[\d+\]', '', teks) 
    teks = re.sub(r'\(\w+ et al\., \d{4}\)', '', teks)
    return teks.strip()

def format_detik_ke_jam(seconds):
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    if h > 0:
        return f"{h:02}:{m:02}:{s:02}"
    return f"{m:02}:{s:02}"

@st.cache_data(show_spinner=False)
def get_audio_slice(audio_bytes, start_s, end_s):
    try:
        from pydub import AudioSegment
        import io
        audio_io = io.BytesIO(audio_bytes)
        audio = AudioSegment.from_file(audio_io)
        sliced = audio[int(start_s * 1000) : int(end_s * 1000)]
        out_io = io.BytesIO()
        sliced.export(out_io, format="mp3")
        return out_io.getvalue()
    except ImportError:
        return None
    except Exception:
        return None
    
# ==========================================
# 4. INIT SESSION STATE & PEMISAHAN TAB UTAMA
# ==========================================
if 'local_files' not in st.session_state: st.session_state.local_files = {}
if 'summary_results' not in st.session_state: st.session_state.summary_results = {}
if 'sub_corpora' not in st.session_state: st.session_state.sub_corpora = {"General": []}
if "uploader_key" not in st.session_state: st.session_state.uploader_key = 0

st.markdown("<h3 style='color:#0F172A;'>📁 Analisis Dokumen Eksternal</h3>", unsafe_allow_html=True)

tab_induk_doc, tab_induk_voice = st.tabs(["📄 Upload Dokumen", "🎙️ Upload Voice (Transkrip)"])


@st.fragment
def render_tab_compare(docs_terpilih, suffix=""):
    st.markdown("<h3 style='color:#0F172A;'>⚖️ Perbandingan Analisis Dokumen (Multi-File)</h3>", unsafe_allow_html=True)
    
    with st.expander("📖 Panduan Membaca Analisis Komparasi", expanded=False):
        st.markdown("""
        Fitur ini membandingkan isi hingga 5 dokumen sekaligus.
        * **📊 Jaccard Index:** Menghitung kosakata yang persis sama dan muncul di *seluruh* dokumen yang dipilih.
        * **🧠 Cosine Similarity:** Menghitung rata-rata kemiripan makna antar-pasangan dokumen.
        * **Tips:** Gunakan Tab di bagian bawah untuk melihat kata eksklusif masing-masing dokumen secara terpisah.
        """)
    
    opsi_semua_doc = docs_terpilih

    if len(opsi_semua_doc) < 2:
        st.info("ℹ️ Tambahkan minimal 2 file dari Sidebar untuk menggunakan fitur Komparasi.")
        return

    with st.container(border=True):
        docs_komparasi = st.multiselect(
            "Pilih 2 hingga 5 file untuk dibandingkan:", 
            opsi_semua_doc, 
            default=opsi_semua_doc[:2] if len(opsi_semua_doc) >= 2 else None,
            max_selections=5,
            key=f"comp_multi_{suffix}"
        )
        
        if len(docs_komparasi) < 2:
            st.warning("⚠️ Silakan pilih minimal 2 file yang berbeda.")
        else:
            if st.button("🚀 Mulai Analisis Komparasi AI", type="primary", use_container_width=True, key=f"btn_comp_multi_{suffix}"):
                st.session_state[f'tampilkan_komp_{suffix}'] = docs_komparasi
                
            if st.session_state.get(f'tampilkan_komp_{suffix}') == docs_komparasi:
                vocabs = [st.session_state.local_files[d]['vocab'] for d in docs_komparasi]
                teks_list = [st.session_state.local_files[d]['cleaned'] for d in docs_komparasi]
                
                irisan_global = set.intersection(*vocabs)
                gabungan_global = set.union(*vocabs)
                jaccard_sim = len(irisan_global) / len(gabungan_global) if len(gabungan_global) > 0 else 0
                
                vectorizer = TfidfVectorizer(stop_words='english')
                tfidf_matrix = vectorizer.fit_transform(teks_list)
                cos_matrix = cosine_similarity(tfidf_matrix)
                
                n = len(docs_komparasi)
                sum_cos = sum(cos_matrix[i][j] for i in range(n) for j in range(i+1, n))
                count_cos = n * (n - 1) / 2
                avg_cosine_sim = sum_cos / count_cos if count_cos > 0 else 0
                
                unik_per_doc = {}
                raw_per_doc = {}
                count_unik_per_doc = {}
                
                nlp_komparasi = load_ai_model(st.session_state.local_files[docs_komparasi[0]].get('lang', 'en'))
                stop_words = nlp_komparasi.Defaults.stop_words
                
                semua_raw = []
                for i, d in enumerate(docs_komparasi):
                    gabungan_lain = set.union(*(vocabs[:i] + vocabs[i+1:]))
                    unik_per_doc[d] = vocabs[i] - gabungan_lain
                    
                    raw_words = re.findall(r'\b[a-z]+(?:-[a-z]+)*\b', teks_list[i].lower())
                    raw_per_doc[d] = raw_words
                    semua_raw.extend(raw_words)
                    
                    count_unik_per_doc[d] = Counter([w for w in raw_words if w in unik_per_doc[d] and w not in stop_words and len(w) > 2])
                
                count_irisan_global = Counter([w for w in semua_raw if w in irisan_global and w not in stop_words and len(w) > 2])

                st.markdown("---")
                col_score1, col_score2 = st.columns(2)
                with col_score1:
                    st.markdown(f"""
                    <div style='background:#F0F9FF; padding:20px; border-radius:12px; border:1px solid #BAE6FD; text-align:center; height:100%;'>
                        <div style='font-size:15px; color:#0369A1; font-weight:700; margin-bottom:5px;'>📊 Jaccard Index Global</div>
                        <div style='font-size:13px; color:#475569; margin-bottom:10px;'>(Kemiripan di seluruh {n} file)</div>
                        <div style='font-size:42px; color:#0EA5E9; font-weight:900;'>{jaccard_sim*100:.1f}%</div>
                    </div>
                    """, unsafe_allow_html=True)
                with col_score2:
                    st.markdown(f"""
                    <div style='background:#FAF5FF; padding:20px; border-radius:12px; border:1px solid #E9D5FF; text-align:center; height:100%; box-shadow: 0 4px 6px -1px rgba(168, 85, 247, 0.1);'>
                        <div style='font-size:15px; color:#6B21A8; font-weight:700; margin-bottom:5px;'>🧠 Avg Cosine Similarity</div>
                        <div style='font-size:13px; color:#475569; margin-bottom:10px;'>(Rata-rata kemiripan makna antar file)</div>
                        <div style='font-size:42px; color:#A855F7; font-weight:900;'>{avg_cosine_sim*100:.1f}%</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                st.markdown("""
                <style>
                div[data-testid="stVerticalBlockBorderWrapper"]:has(span[title="marker-biru"]) { background-color: #F0F9FF !important; border-color: #BAE6FD !important; }
                div[data-testid="stVerticalBlockBorderWrapper"]:has(span[title="marker-biru"]) [data-testid="stPills"] button { background-color: #FFFFFF !important; border: 1px solid #CBD5E1 !important; }
                div[data-testid="stVerticalBlockBorderWrapper"]:has(span[title="marker-biru"]) [data-testid="stPills"] button p,
                div[data-testid="stVerticalBlockBorderWrapper"]:has(span[title="marker-biru"]) [data-testid="stPills"] button span { color: #0284C7 !important; font-weight: 600 !important; }
                div[data-testid="stVerticalBlockBorderWrapper"]:has(span[title="marker-biru"]) [data-testid="stPills"] button:hover { border-color: #0EA5E9 !important; background-color: #F0F9FF !important; }
                </style>
                """, unsafe_allow_html=True)

                def ekstrak_kalimat(kata, docs_to_search):
                    kalimat_terkumpul = []
                    pola = re.compile(r'\b' + re.escape(kata) + r'\b', re.IGNORECASE)
                    for doc_name in docs_to_search:
                        teks = st.session_state.local_files[doc_name]['cleaned']
                        sents = re.split(r'(?<=[.!?]) +', teks)
                        for s in sents:
                            if pola.search(s):
                                s_aman = s.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("$", "&#36;").replace("{", "&#123;").replace("}", "&#125;").replace("`", "&#96;")
                                s_high = pola.sub(r"<mark style='background:#FDE047; padding:2px 6px; border-radius:4px; font-weight:bold; color:#0F172A;'>\g<0></mark>", s_aman)
                                kalimat_terkumpul.append({"doc": doc_name, "text": s_high})
                    return kalimat_terkumpul

                def render_pills_box(judul, count_data, marker_id, icon, prefix):
                    top_words = count_data.most_common(30)
                    options = [f"{w} ({freq}x)" for w, freq in top_words]
                    with st.container(border=True):
                        st.markdown(f"<span title='{marker_id}'></span><h4 style='margin:0 0 10px 0; color:#0F172A;'>{icon} {judul}</h4>", unsafe_allow_html=True)
                        if not options:
                            st.caption("Tidak ada kata eksklusif/signifikan yang membedakan dokumen ini.")
                            return None
                        else:
                            try:
                                return st.pills("Pilih Kata", options, label_visibility="collapsed", key=f"pills_{prefix}_{suffix}")
                            except AttributeError:
                                return st.selectbox("Pilih Kata:", ["-- Pilih --"] + options, key=f"pills_{prefix}_{suffix}")

                def aksi_prev_kalimat(state_key): st.session_state[state_key] -= 1
                def aksi_next_kalimat(state_key): st.session_state[state_key] += 1
                def aksi_loncat_kalimat(state_key, input_key): st.session_state[state_key] = st.session_state[input_key] - 1

                def render_awan_konteks_inline(active_pill, active_docs, prefix_key):
                    nama_state_word = f"ctx_word_{prefix_key}_{suffix}"
                    nama_state_page = f"ctx_page_{prefix_key}_{suffix}"
                    nama_state_sents = f"ctx_sents_{prefix_key}_{suffix}"
                    
                    if st.session_state.get(nama_state_word) != active_pill or nama_state_sents not in st.session_state:
                        st.session_state[nama_state_page] = 0
                        st.session_state[nama_state_word] = active_pill
                        st.session_state[nama_state_sents] = ekstrak_kalimat(active_pill, active_docs)
                        
                    sents = st.session_state.get(nama_state_sents, [])
                    tot = len(sents)
                    if tot == 0: return
                        
                    p = st.session_state.get(nama_state_page, 0)
                    st.markdown(
                        f"<div style='background-color: #ffffff; border: 1px dashed #0EA5E9; border-radius: 8px; padding: 15px; margin-top: 10px; margin-bottom: 15px;'>"
                        f"<h5 style='color:#0EA5E9; margin:0 0 8px 0; font-size:14px; border-bottom: 1px solid #f1f5f9; padding-bottom:5px;'>☁️ Konteks: <b style='color:#0F172A;'>{active_pill}</b></h5>"
                        f"<div style='color: #64748b; font-size: 13px;'>📄 {sents[p]['doc']}</div>"
                        f"<div style='color: #0f172a; font-size: 13.5px; line-height: 1.5; margin-top:5px;'>\"{sents[p]['text']}\"</div>"
                        f"</div>", unsafe_allow_html=True
                    )
                    
                    c_p, c_i, c_n = st.columns([2, 3, 2])
                    c_p.button("⬅️", disabled=(p == 0), key=f"prev_{prefix_key}_{suffix}", use_container_width=True, on_click=aksi_prev_kalimat, args=(nama_state_page,))
                    with c_i:
                        with st.popover(f"{p+1} / {tot}", use_container_width=True):
                            st.number_input("Lompat ke:", min_value=1, max_value=max(1, tot), value=p+1, key=f"num_{prefix_key}_{suffix}", on_change=aksi_loncat_kalimat, args=(nama_state_page, f"num_{prefix_key}_{suffix}"))
                    c_n.button("➡️", disabled=(p == tot - 1), key=f"next_{prefix_key}_{suffix}", use_container_width=True, on_click=aksi_next_kalimat, args=(nama_state_page,))
                def render_panel(col, sents, p, p_key, doc_name, color_hex, bg_light):
                    with col:
                        tot = len(sents)
                        if tot == 0:
                            st.info("Konteks tidak ditemukan.")
                            return
                        st.markdown(
                            f"<div style='background-color: {bg_light}; border: 1px solid {color_hex}; border-radius: 8px; padding: 15px; margin-bottom: 10px; height: 100%;'>"
                            f"<h6 style='color:{color_hex}; margin:0 0 8px 0; font-size:13px; border-bottom: 1px solid {color_hex}44; padding-bottom:5px;'>📄 {doc_name}</h6>"
                            f"<div style='color: #0f172a; font-size: 13.5px; line-height: 1.6; min-height: 80px;'>\"{sents[p]['text']}\"</div>"
                            f"</div>", unsafe_allow_html=True
                        )
                        c_p, c_i, c_n = st.columns([2, 3, 2])
                        c_p.button("⬅️", disabled=(p == 0), key=f"p_{p_key}", use_container_width=True, on_click=aksi_prev_kalimat, args=(p_key,))
                        with c_i:
                            with st.popover(f"{p+1} / {tot}", use_container_width=True):
                                st.number_input("Lompat ke:", min_value=1, max_value=max(1, tot), value=p+1, key=f"num_{p_key}", on_change=aksi_loncat_kalimat, args=(p_key, f"num_{p_key}"))                        
                        c_n.button("➡️", disabled=(p == tot - 1), key=f"n_{p_key}", use_container_width=True, on_click=aksi_next_kalimat, args=(p_key,))

                st.write("")
                st.markdown("<h4 style='color:#0F172A;'>💎 Kata Eksklusif per Dokumen</h4>", unsafe_allow_html=True)
                tabs = st.tabs([f"📄 File {i+1}" for i in range(n)])
                
                for i, d in enumerate(docs_komparasi):
                    with tabs[i]:
                        st.caption(f"Kata yang HANYA muncul di **{d}** dan tidak ada di file lainnya.")
                        pil_unik = render_pills_box(f"Eksklusif: {d}", count_unik_per_doc[d], "marker-biru", "📘", f"eks_{i}")
                        if pil_unik and pil_unik != "-- Pilih --":
                            render_awan_konteks_inline(pil_unik.split(" ")[0], [d], f"aw_eks_{i}")

                st.write("")
                st.markdown("---")
                pil_iris = render_pills_box(f"🤝 Kata Beririsan (Muncul di seluruh {n} file)", count_irisan_global, "marker-biru", "🔗", "iris_global")
                
                if pil_iris and pil_iris != "-- Pilih --":
                    kata_iris = pil_iris.split(" ")[0]
                    st.markdown(f"<div style='margin-top:15px; margin-bottom:5px; padding-bottom:5px; border-bottom:2px solid #0EA5E9;'><h4 style='color:#0F172A; margin:0;'>🔍 Konteks Bersama: <b style='color:#0284C7;'>{kata_iris}</b></h4></div>", unsafe_allow_html=True)
                    
                    if st.session_state.get(f"ctx_w_iris_{suffix}") != kata_iris:
                        st.session_state[f"ctx_w_iris_{suffix}"] = kata_iris
                        for d in docs_komparasi:
                            st.session_state[f"ctx_s_{d}_iris_{suffix}"] = ekstrak_kalimat(kata_iris, [d])
                            st.session_state[f"ctx_p_{d}_iris_{suffix}"] = 0
                            
                    cols = st.columns(n)
                    colors = ["#0EA5E9", "#E11D48", "#D97706", "#7C3AED", "#059669"]
                    bgs = ["#F0F9FF", "#FDF4FF", "#FFFBEB", "#F5F3FF", "#ECFDF5"]
                    
                    for i, d in enumerate(docs_komparasi):
                        sents = st.session_state.get(f"ctx_s_{d}_iris_{suffix}", [])
                        p = st.session_state.get(f"ctx_p_{d}_iris_{suffix}", 0)
                        p_key = f"ctx_p_{d}_iris_{suffix}"
                        render_panel(cols[i], sents, p, p_key, d, colors[i%5], bgs[i%5])

                st.markdown("<div style='height: 30px; display: block;'></div>", unsafe_allow_html=True)
                            
                            
@st.fragment
def render_tab_visual(docs_terpilih, suffix=""):
    st.markdown("<h3 style='color:#0F172A;'>📈 Visualisasi Detail</h3>", unsafe_allow_html=True)
    st.markdown("""
        <style>
        .scrollable-column { max-height: 800px; overflow-y: auto; padding-right: 15px; scrollbar-width: thin; }
        .scrollable-column::-webkit-scrollbar { width: 6px; }
        .scrollable-column::-webkit-scrollbar-track { background: #F1F5F9; border-radius: 4px; }
        .scrollable-column::-webkit-scrollbar-thumb { background: #CBD5E1; border-radius: 4px; }
        .scrollable-column::-webkit-scrollbar-thumb:hover { background: #94A3B8; }
        </style>
    """, unsafe_allow_html=True)

    if len(docs_terpilih) == 0:
        st.warning("Silakan pilih file terlebih dahulu.")
        return

    docs_to_visualize = st.multiselect(
        "Pilih maksimal 2 file untuk divisualisasikan berdampingan:", 
        docs_terpilih, 
        default=docs_terpilih[:2] if len(docs_terpilih) >= 2 else docs_terpilih, 
        max_selections=2, 
        key=f"vis_sel_{suffix}"
    )

    if not docs_to_visualize:
        return

    state_key = f"is_rendered_{suffix}"
    
    if state_key not in st.session_state:
        st.session_state[state_key] = False

    if st.button("🚀 Render Visualisasi", type="primary", key=f"btn_render_{suffix}"):
        st.session_state[state_key] = True

    if not st.session_state[state_key]:
        st.info("Klik tombol 'Render Visualisasi' di atas untuk mulai memproses data.")
        return

    def render_visualisasi_sisi(fname, warna_utama, warna_kedua):
        st.markdown(f"<div style='background:{warna_utama}22; padding:10px 15px; border-radius:8px; border-bottom:3px solid {warna_utama}; margin-bottom:15px;'><h4 style='margin:0; color:#1E293B;'>📄 {fname}</h4></div>", unsafe_allow_html=True)
        st.markdown("<div class='scrollable-column'>", unsafe_allow_html=True)
        
        if 'vis_cache' not in st.session_state.local_files[fname]:
            with st.spinner(f"Memproses visualisasi {fname}..."):
                teks_dokumen = st.session_state.local_files[fname]['text']
                teks_mentah_aktif = st.session_state.local_files[fname]['cleaned']
                nlp_aktif = load_ai_model(st.session_state.local_files[fname].get('lang', 'en'))
                df_pos, df_words, df_cloud_text = dapatkan_data_visual(teks_dokumen[:20000], nlp_aktif)
                
                if not df_words.empty and 'Pasangan 1' not in df_words.columns:
                    col1, col2, col3, col4, col5 = [], [], [], [], []
                    for i, kata in enumerate(df_words['Kata']):
                        if i < 500:  
                            hasil_colloc = hitung_collocation(kata, teks_mentah_aktif, nlp_aktif, window=5)
                            for col, idx in zip([col1, col2, col3, col4, col5], range(5)):
                                col.append(f"{hasil_colloc[idx][0]} ({hasil_colloc[idx][1]}x)" if len(hasil_colloc) > idx else "-")
                        else:
                            for col in [col1, col2, col3, col4, col5]: col.append("-")
                    df_words['Pasangan 1'] = col1; df_words['Pasangan 2'] = col2; df_words['Pasangan 3'] = col3; df_words['Pasangan 4'] = col4; df_words['Pasangan 5'] = col5

                words_ng = [w for w in re.findall(r'\b[a-z]{3,}\b', teks_mentah_aktif.lower()) if w not in nlp_aktif.Defaults.stop_words]
                bigrams = [" ".join(g) for g in zip(*[words_ng[i:] for i in range(2)])]
                trigrams = [" ".join(g) for g in zip(*[words_ng[i:] for i in range(3)])]
                df_bigram = pd.DataFrame(Counter(bigrams).most_common(15), columns=['Frasa', 'Frekuensi'])
                df_trigram = pd.DataFrame(Counter(trigrams).most_common(15), columns=['Frasa', 'Frekuensi'])

                cloud_img = None
                if df_cloud_text:
                    fig = get_cached_wordcloud(df_cloud_text[:20000])
                    buf = io.BytesIO(); fig.savefig(buf, format="png", bbox_inches='tight', pad_inches=0); plt.close(fig)
                    cloud_img = buf.getvalue()

                st.session_state.local_files[fname]['vis_cache'] = {'df_pos': df_pos, 'df_words': df_words, 'df_bigram': df_bigram, 'df_trigram': df_trigram, 'cloud_img': cloud_img}

        cache_vis = st.session_state.local_files[fname]['vis_cache']
        st.markdown(f"**☁️ Word Cloud**")
        if cache_vis['cloud_img']: st.image(cache_vis['cloud_img'], use_container_width=True)
        st.write("---")
        st.markdown(f"**📈 15 Kata Teratas**")
        if not cache_vis['df_words'].empty: 
            st.altair_chart(alt.Chart(cache_vis['df_words'].head(15)).mark_bar(color=warna_utama, cornerRadiusEnd=4).encode(y=alt.Y('Kata:N', sort='-x', title=None), x=alt.X('Frekuensi:Q', title=None), tooltip=[alt.Tooltip('Kata:N'), alt.Tooltip('Frekuensi:Q'), alt.Tooltip('Pasangan 1:N', title='🔗 Colloc 1')]).properties(height=350), use_container_width=True)
        st.write("---")
        st.markdown(f"**🔗 15 Frasa Bigram (2 Kata)**")
        if not cache_vis['df_bigram'].empty: 
            st.altair_chart(alt.Chart(cache_vis['df_bigram']).mark_bar(color=warna_kedua, cornerRadiusEnd=4).encode(y=alt.Y('Frasa:N', sort='-x', title=None), x=alt.X('Frekuensi:Q', title=None), tooltip=['Frasa', 'Frekuensi']).properties(height=350), use_container_width=True)
        st.write("---")
        st.markdown(f"**🔗 15 Frasa Trigram (3 Kata)**")
        if not cache_vis['df_trigram'].empty: 
            st.altair_chart(alt.Chart(cache_vis['df_trigram']).mark_bar(color=warna_utama, opacity=0.8, cornerRadiusEnd=4).encode(y=alt.Y('Frasa:N', sort='-x', title=None), x=alt.X('Frekuensi:Q', title=None), tooltip=['Frasa', 'Frekuensi']).properties(height=350), use_container_width=True)
        st.write("---")
        st.markdown(f"**🔤 Statistik Tata Bahasa (POS)**")
        if not cache_vis['df_pos'].empty: 
            st.altair_chart(alt.Chart(cache_vis['df_pos']).mark_bar(color=warna_kedua, opacity=0.8, cornerRadiusEnd=4).encode(y=alt.Y('POS Tag', sort='-x', title=None), x=alt.X('Jumlah Kata', title=None), tooltip=['POS Tag', 'Jumlah Kata']).properties(height=350), use_container_width=True)
        if not cache_vis['df_words'].empty:
            st.write("---")
            st.markdown(f"**🗃️ Tabel Data Collocation**")
            df_tabel_500 = cache_vis['df_words'].head(500)
            st.dataframe(df_tabel_500, use_container_width=True, height=250, hide_index=True)
            buffer_excel = io.BytesIO()
            with pd.ExcelWriter(buffer_excel, engine='openpyxl') as writer:
                df_tabel_500.to_excel(writer, index=False, sheet_name='Concordance')
            
            st.download_button(
                label="📥 Unduh Data Excel (XLSX)", 
                data=buffer_excel.getvalue(), 
                file_name=f"concordance_{fname}.xlsx", 
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                key=f"dl_excel_{fname}_{suffix}"
            )            
        st.markdown("</div>", unsafe_allow_html=True)

    if len(docs_to_visualize) == 1:
        with st.container(border=True): render_visualisasi_sisi(docs_to_visualize[0], "#0EA5E9", "#38BDF8")
    elif len(docs_to_visualize) == 2:
        with st.container(border=True):
            col_vis_kiri, col_vis_kanan = st.columns(2)
            with col_vis_kiri: render_visualisasi_sisi(docs_to_visualize[0], "#0EA5E9", "#38BDF8")
            with col_vis_kanan: render_visualisasi_sisi(docs_to_visualize[1], "#D946EF", "#E879F9")

@st.fragment
def render_tab_search(docs_terpilih):
    def prev_page_1(): st.session_state.current_page -= 1
    def next_page_1(): st.session_state.current_page += 1
    def aksi_loncat_halaman_search(): st.session_state.current_page = st.session_state.num_search_page - 1
    def update_ws_word(new_word): st.session_state["in_word_sketch"] = new_word 

    st.markdown("<h3 style='color:#0F172A;'>🔍 Pencarian Pintar & Pemrofilan Kata</h3>", unsafe_allow_html=True)
    with st.expander("📖 Panduan Mode Pencarian AI (Untuk Pemula)", expanded=False):
        st.markdown("""
        **Gunakan kecerdasan buatan (AI) untuk menemukan informasi spesifik dari tumpukan teks:**
        
        * 🔍 **Lemmatization:** Mencari kata dasar beserta semua variasinya. *(Ketik "lari", maka "berlari" dan "melarikan" juga akan tersorot otomatis).*
        * 🔗 **Collocation (Teman Sanding):** Menemukan kata apa yang paling sering muncul *berdekatan* dengan kata kunci Anda. Semakin tinggi skor PMI, semakin erat hubungannya di dalam teks.
        * 🔬 **Morphology:** Membedah satu kata dasar menjadi grafik variasi imbuhan yang digunakan dalam dokumen.
        * 🧠 **Semantic Search:** Mencari kalimat yang *maknanya mirip* atau sejalan dengan apa yang Anda ketik, meskipun susunan katanya sama sekali berbeda.
        * 🧩 **Word Sketch:** Memprofilkan suatu kata. Melihat kata sifat apa yang sering menggambarkannya, atau tindakan (kata kerja) apa yang sering ia lakukan.
        * 🏷️ **Entity Search (NER):** Mesin otomatis melacak dan menyorot entitas spesifik seperti **Nama Orang, Organisasi, Lokasi, atau Nominal Uang** tanpa perlu Anda ketik satu per satu.
        * 📚 **POS Search:** Mencari berdasarkan *jabatan kata* (Misal: temukan semua 'Kata Sifat' di dokumen ini).
        * ⚙️ **Fitur Lanjutan:** Gunakan **Regex** untuk mencari format pola (seperti email/tanggal), **Boolean** untuk logika (Kata A *DAN* Kata B), atau **Dependency** untuk melihat struktur anak-induk kalimat.
        """)

    col_mode, col_input, col_btn = st.columns([2.5, 4.5, 1], gap="small")
    with col_mode:
        mode_pencarian = st.selectbox("Mode Pencarian", [
            "🔍 Lemmatization", "🔗 Collocation (Asosiasi Korpus)", "🔬 Morphology Search (Variasi Kata)", "🧠 Semantic Search", "🧩 Word Sketch (Profil Kata)", 
            "🏷️ Entity Search (NER)", "📚 POS Search (Kelas Kata)", "🛒 Boolean Search", "⚙️ Regex Search", "🌳 Dependency Search"
        ], label_visibility="collapsed")
        
    if 'last_search_mode' not in st.session_state:
        st.session_state.last_search_mode = mode_pencarian

    if st.session_state.last_search_mode != mode_pencarian:
        st.session_state.last_search_mode = mode_pencarian
        st.session_state.last_query = "" 
        
        st.session_state.search_results = []
        st.session_state.colloc_results = pd.DataFrame()
        st.session_state.morph_results = pd.DataFrame()
        st.session_state.sketch_results = {}
        st.session_state.current_page = 0
        
    with col_input:
        if "NER" in mode_pencarian:
            query_aktif = st.selectbox("Pilih Tipe Entitas", ["PERSON (Orang)", "ORG (Organisasi)", "GPE (Negara/Kota)", "LOC (Lokasi)", "DATE (Tanggal)", "MONEY (Keuangan)"], key="input_search_ner", label_visibility="collapsed")
        elif "POS Search" in mode_pencarian:
            c_pos1, c_pos2 = st.columns([1, 1])
            with c_pos1: pos_label = st.selectbox("Pilih Kelas Kata:", list(MAP_SEMUA_POS.keys()), key="pos_tag_sel", label_visibility="collapsed"); target_pos_tag = MAP_SEMUA_POS[pos_label]
            with c_pos2: pos_keyword = st.text_input("Kata Kunci (Opsional):", placeholder="Ketik kata...", key="pos_kw_in", label_visibility="collapsed").strip().lower()
            query_aktif = f"POS_{target_pos_tag}_{pos_keyword}"
        elif "Dependency" in mode_pencarian:
            c_dep1, c_dep2, c_dep3 = st.columns(3)
            with c_dep1: head_word = st.text_input("Head (Induk):", placeholder="e.g. train", key="dep_h").strip().lower()
            with c_dep2: rel_type = st.selectbox("Relation:", ["nsubj", "obj", "dobj", "amod", "advmod", "compound", "prep"], key="dep_r")
            with c_dep3: child_word = st.text_input("Child (Anak):", placeholder="e.g. model", key="dep_c").strip().lower()
            query_aktif = f"{head_word} {rel_type} {child_word}".strip()
        elif "Regex" in mode_pencarian:
            opsi_regex = {
                "📖 Sitasi (Penulis, Tahun)": r"\([A-Z][A-Za-z\s]+(?:et al\.)?,\s?\d{4}\)",
                "🔖 Referensi Gambar/Tabel": r"\b(?:Gambar|Figure|Tabel|Table)\s+\d+(?:\.\d+)*\b",
                "📧 Format Email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
                "🔗 Link URL/Website": r"https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+",
                "🧪 Persamaan Matematika": r"([a-zA-Z\d\s\+\-\*\/\=\^\(\)]+\s?=\s?[a-zA-Z\d\s\+\-\*\/\^\(\)]+)",
                "📑 Format DOI": r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b",
                "📅 Format Tanggal (DD-MM-YYYY)": r"\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b",
                "🔢 Angka Persentase": r"\b\d+(?:\.\d+)?%",
                "💵 Nilai Mata Uang (IDR/USD)": r"\b(?:Rp|IDR|USD|\$)\s?\d{1,3}(?:\.\d{3})*(?:,\d{2})?\b",
                "🕒 Format Waktu (HH:MM)": r"\b(?:[01]?\d|2[0-3]):[0-5]\d(?::[0-5]\d)?\b",
                "🏠 Alamat IP (IPv4)": r"\b(?:\d{1,3}\.){3}\d{1,3}\b",
                "📚 Judul BAB (Kapital)": r"^BAB\s+[IVXLCDM]+\b.*",
                "🏛️ Kode Institusi/Negara": r"\b[A-Z]{3,}\b",
                "✏️ Ketik Manual...": "manual"
            }     
            pilihan_regex = st.selectbox("Pilih Pola Regex", list(opsi_regex.keys()), label_visibility="collapsed")
            if pilihan_regex == "✏️ Ketik Manual...": query_aktif = st.text_input("Ketik pola Regex", placeholder="Contoh: \\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\\.[A-Z]{2,}\\b", label_visibility="collapsed")
            else: query_aktif = opsi_regex[pilihan_regex]
            
        elif "Collocation" in mode_pencarian:
            c_col1, c_col2 = st.columns([3, 1])
            with c_col1: query_aktif = st.text_input("Ketik Kata Pusat (Node):", placeholder="Ketik kata...", key="in_colloc", label_visibility="collapsed").strip().lower()
            with c_col2: window_size = st.selectbox("Rentang Jarak (±)", [3, 4, 5, 8, 10], index=1, label_visibility="collapsed")

        elif "Word Sketch" in mode_pencarian: query_aktif = st.text_input("Ketik kata untuk diprofilkan:", placeholder="Contoh: model, data...", key="in_word_sketch", label_visibility="collapsed").strip().lower()
        elif "Semantic" in mode_pencarian: query_aktif = st.text_input("Ketik konsep makna", placeholder="Ketik kata...", key="in_semantic", label_visibility="collapsed").strip()
        elif "Lemmatization" in mode_pencarian: query_aktif = st.text_input("Ketik kata dasar", placeholder="Ketik kata...", key="in_lemma", label_visibility="collapsed").strip()
        elif "Morphology" in mode_pencarian: query_aktif = st.text_input("Ketik Akar Kata", placeholder="Ketik kata...", key="in_morph", label_visibility="collapsed").strip()
        elif "Boolean" in mode_pencarian: query_aktif = st.text_input("Query Boolean", placeholder="Contoh: machine AND translation NOT neural", key="in_bool", label_visibility="collapsed")

    with col_btn:
        btn_cari = st.button("Cari", key="btn_search_key", use_container_width=True, type="primary")

    if 'search_results' not in st.session_state: st.session_state.search_results = []
    if 'morph_results' not in st.session_state: st.session_state.morph_results = pd.DataFrame()
    if 'sketch_results' not in st.session_state: st.session_state.sketch_results = {}
    if 'colloc_results' not in st.session_state: st.session_state.colloc_results = pd.DataFrame()
    if 'current_page' not in st.session_state: st.session_state.current_page = 0
    if 'last_query' not in st.session_state: st.session_state.last_query = ""

    def aksi_klik_saran():
        pilihan = st.session_state.get('saran_pills_widget')
        if pilihan:
            st.session_state.trigger_search = True
            st.session_state.current_page = 0
            if "Word Sketch" in mode_pencarian: st.session_state["in_word_sketch"] = pilihan
            elif "Semantic" in mode_pencarian: st.session_state["in_semantic"] = pilihan
            elif "Lemmatization" in mode_pencarian: st.session_state["in_lemma"] = pilihan
            elif "Morphology" in mode_pencarian: st.session_state["in_morph"] = pilihan
            elif "Collocation" in mode_pencarian: st.session_state["in_colloc"] = pilihan
            else: st.session_state["teks_pencarian"] = pilihan

    if query_aktif:
        if st.session_state.get('last_query') != query_aktif or btn_cari or st.session_state.get('trigger_search'):
            st.session_state.trigger_search = False
            lang_query = st.session_state.local_files[docs_terpilih[0]]['lang'] if docs_terpilih else 'en'
            nlp_query = load_ai_model(lang_query)
            
            if "Collocation" in mode_pencarian:
                try:
                    node_word = query_aktif.lower().strip()
                    freq_node = 0
                    freq_collocate = Counter()
                    freq_cooc = Counter()
                    total_words = 0
                    
                    for fname in docs_terpilih:
                        df_indeks = st.session_state.local_files[fname].get('index_df')
                        
                        if df_indeks is None or df_indeks.empty:
                            st.warning(f"⚠️ Cache Indeks file '{fname}' kosong. Silakan hapus dan upload ulang file ini.")
                            continue
                            
                        tokens = []
                        for t_list, p_list in zip(df_indeks['Tokens'], df_indeks['POS_Tags']):
                            for t, p in zip(t_list, p_list):
                                if p not in ['SPACE', 'PUNCT', 'SYM']: 
                                    tokens.append(t)
                                    
                        total_words += len(tokens)
                        
                        for i, token in enumerate(tokens):
                            freq_collocate[token] += 1
                            if token == node_word:
                                freq_node += 1
                                start = max(0, i - window_size)
                                end = min(len(tokens), i + window_size + 1)
                                for j in range(start, end):
                                    if i != j: 
                                        freq_cooc[tokens[j]] += 1
                    
                    colloc_data = []
                    if freq_node > 0:
                        for col_word, f_co in freq_cooc.items():
                            if f_co >= 2: 
                                f_col = freq_collocate[col_word]
                                pmi = math.log2((f_co * total_words) / (freq_node * f_col))
                                colloc_data.append({
                                    "Kata Sanding (Collocate)": col_word,
                                    "Frekuensi Bersama": f_co,
                                    "Frekuensi Total (Korpus)": f_col,
                                    "Skor Asosiasi (PMI)": round(pmi, 3)
                                })
                        
                        df_colloc = pd.DataFrame(colloc_data)
                        if not df_colloc.empty:
                            df_colloc = df_colloc.sort_values(by="Skor Asosiasi (PMI)", ascending=False).reset_index(drop=True)
                        st.session_state.colloc_results = df_colloc
                    else:
                        st.session_state.colloc_results = pd.DataFrame()
                        
                    st.session_state.search_results, st.session_state.sketch_results, st.session_state.morph_results = [], {}, pd.DataFrame()
                    st.session_state.last_query = query_aktif
                except Exception as e: 
                    st.error(f"Error Collocation: {e}")

            elif "Morphology" in mode_pencarian:
                try:
                    target_lemma = nlp_query(query_aktif)[0].lemma_.lower()
                    morph_data = []
                    for fname in docs_terpilih: 
                        teks_mentah = str(st.session_state.local_files[fname]['cleaned'])
                        doc = get_cached_spacy_doc(teks_mentah[:20000], st.session_state.local_files[fname]['lang'])
                        for token in doc:
                            if token.lemma_.lower() == target_lemma and not token.is_punct and not token.is_space:
                                morph_data.append({"Bentuk Teks Asli": token.text.lower(), "Kelas Kata (POS)": token.pos_, "Struktur Morfologi": str(token.morph) if str(token.morph) else "Bentuk Dasar (Base)"})
                    if morph_data:
                        st.session_state.morph_results = pd.DataFrame(morph_data).groupby(['Bentuk Teks Asli', 'Kelas Kata (POS)', 'Struktur Morfologi']).size().reset_index(name='Frekuensi Muncul').sort_values('Frekuensi Muncul', ascending=False).reset_index(drop=True)
                    else: st.session_state.morph_results = pd.DataFrame()
                    st.session_state.search_results, st.session_state.sketch_results, st.session_state.colloc_results, st.session_state.last_query = [], {}, pd.DataFrame(), query_aktif
                except Exception as e: st.error(f"Error: {e}")
                    
            elif "Word Sketch" in mode_pencarian:
                try:
                    st.session_state.sketch_results = buat_word_sketch(query_aktif, " ".join([st.session_state.local_files[f]['cleaned'] for f in docs_terpilih]), nlp_query)
                    st.session_state.search_results, st.session_state.morph_results, st.session_state.colloc_results, st.session_state.last_query = [], pd.DataFrame(), pd.DataFrame(), query_aktif
                except Exception as e: st.error(f"Error: {e}")
            
            else:
                matches_global = []
                if "Lemmatization" in mode_pencarian or "Semantic" in mode_pencarian: query_doc = nlp_query(query_aktif)
                
                for fname in docs_terpilih: 
                    doc = get_cached_spacy_doc(st.session_state.local_files[fname]['cleaned'], st.session_state.local_files[fname]['lang'])
                    
                    if "POS Search" in mode_pencarian:
                        for s in doc.sents:
                            match_found, highlighted = False, ""
                            for token in s:
                                if (not pos_keyword or token.lemma_.lower() == pos_keyword) and (target_pos_tag == "ALL" or token.pos_ == target_pos_tag) and (pos_keyword or target_pos_tag != "ALL"):
                                    match_found = True; highlighted += f"<mark style='background:#0EA5E9; color:white; font-weight:bold; padding:0 4px; border-radius:4px;'>{token.text}</mark>{token.whitespace_}"
                                else: highlighted += f"{token.text}{token.whitespace_}"
                            if match_found:
                                p_html = "<div style='display:flex; gap:5px; flex-wrap:wrap;'>" + "".join([f"<span style='background-color: {Warna_POS_Utama.get(p, '#94A3B8')}; color:white; border-radius:4px; padding:3px 8px; font-size:11px; font-weight:600;'>{p}: {c}</span>" for p, c in Counter([t.pos_ for t in s if t.pos_ not in ['SPACE', 'PUNCT', 'SYM', 'X']]).items()]) + "</div>"
                                matches_global.append({'file': fname, 'text': s.text.strip(), 'html': highlighted, 'pills': p_html, 'tags': list(set([t.pos_ for t in s if t.pos_ in deskripsi_pos]))})
                                
                    elif "NER" in mode_pencarian:
                        target_ent = query_aktif.split(" ")[0]
                        for s in doc.sents:
                            if any(ent.label_ == target_ent for ent in s.ents):
                                m_text = s.text.strip()
                                p_html = "<div style='display:flex; gap:5px; flex-wrap:wrap;'>" + "".join([f"<span style='background-color: {Warna_POS_Utama.get(p, '#94A3B8')}; color:white; border-radius:4px; padding:3px 8px; font-size:11px; font-weight:600;'>{p}: {c}</span>" for p, c in Counter([t.pos_ for t in s if t.pos_ not in ['SPACE', 'PUNCT', 'SYM', 'X']]).items()]) + "</div>"
                                matches_global.append({'file': fname, 'text': m_text, 'html': get_colored_ner_inline(load_ai_model(st.session_state.local_files[fname]['lang']), m_text, target_ent), 'pills': p_html, 'tags': list(set([t.pos_ for t in s if t.pos_ in deskripsi_pos]))})
                    
                    elif "Dependency" in mode_pencarian:
                        if head_word != "" or child_word != "":
                            for s in doc.sents:
                                match_found, highlighted = False, s.text.strip()
                                for token in s:
                                    if token.dep_ == rel_type or (rel_type == "obj" and token.dep_ == "dobj"):
                                        if (head_word == "" or token.head.lemma_.lower() == head_word) and (child_word == "" or token.lemma_.lower() == child_word):
                                            match_found = True
                                            if head_word != "": highlighted = re.sub(f"\\b({token.head.text})\\b", r"<mark style='background:#F59E0B; color:white; padding:0 4px; border-radius:3px;'>\1</mark>", highlighted, flags=re.I)
                                            if child_word != "": highlighted = re.sub(f"\\b({token.text})\\b", r"<mark style='background:#0EA5E9; color:white; padding:0 4px; border-radius:3px;'>\1</mark>", highlighted, flags=re.I)
                                            break
                                if match_found:
                                    p_html = "<div style='display:flex; gap:5px; flex-wrap:wrap;'>" + "".join([f"<span style='background-color: {Warna_POS_Utama.get(p, '#94A3B8')}; color:white; border-radius:4px; padding:3px 8px; font-size:11px; font-weight:600;'>{p}: {c}</span>" for p, c in Counter([t.pos_ for t in s if t.pos_ not in ['SPACE', 'PUNCT', 'SYM', 'X']]).items()]) + "</div>"
                                    matches_global.append({'file': fname, 'text': s.text.strip(), 'html': highlighted, 'pills': p_html, 'tags': list(set([t.pos_ for t in s if t.pos_ in deskripsi_pos]))})
                    
                    elif "Regex" in mode_pencarian:
                        try:
                            pola_regex = re.compile(query_aktif)
                            sents = nltk.sent_tokenize(st.session_state.local_files[fname]['cleaned'])
                            for m_text in sents:
                                m_text = m_text.strip()
                                if pola_regex.search(m_text):
                                    matches_global.append({'file': fname, 'text': m_text, 'html': pola_regex.sub(r"<mark style='background:#F59E0B; color:white; padding:0 4px; border-radius:3px;'>\g<0></mark>", m_text), 'pills': "", 'tags': []})
                        except: pass

                    elif "Lemmatization" in mode_pencarian:
                        q_l = query_doc[0].lemma_.lower().strip() if len(query_doc) > 0 else query_aktif.lower().strip()
                        df_indeks = st.session_state.local_files[fname].get('index_df')
                        if df_indeks is None or df_indeks.empty:
                            st.warning(f"⚠️ Cache Indeks untuk file '{fname}' belum terbentuk. Silakan HAPUS file ini dari daftar (di bawah) dan UPLOAD ULANG agar fitur pencarian super-cepat bekerja.")
                            continue
                            
                        baris_cocok = df_indeks[df_indeks.apply(lambda row: q_l in row['Lemmas'] or q_l in row['Tokens'] or query_aktif.lower() in row['Tokens'], axis=1)]
                        
                        for _, row in baris_cocok.iterrows():
                            teks_kalimat = row['Teks_Asli']
                            id_segmen = row['ID_Segmen']
                            
                            s = nlp_query(teks_kalimat)
                            
                            matches = [t for t in s if t.lemma_.lower().strip() == q_l or t.text.lower().strip() == q_l or t.text.lower().strip() == query_aktif.lower()]
                            
                            if matches:
                                for t in matches:
                                    idx_in_sent = t.idx
                                    left_raw = s.text[:idx_in_sent].strip()
                                    right_raw = s.text[idx_in_sent + len(t.text):].strip()
                                    if len(left_raw) > 90: 
                                        left_raw = left_raw[-90:].strip()
                                    if len(right_raw) > 90: 
                                        right_raw = right_raw[:90].strip()
                                    left_context = html.escape(left_raw)
                                    keyword = html.escape(t.text)
                                    right_context = html.escape(right_raw)
                                    
                                    h_light_std = s.text.strip()
                                    h_light_std = re.sub(f"\\b({re.escape(t.text)})\\b", r"<mark style='background:#0EA5E9; color:white; padding:0 4px; border-radius:3px;'>\1</mark>", h_light_std, flags=re.I)
                                    
                                    h_light_kwic = (
                                        "<div style='display:flex; justify-content:center; align-items:center; width:100%; "
                                        "font-family:\"Times New Roman\", Times, serif; font-size:16px; padding:10px 0; border-bottom:1px solid #E2E8F0;'>"
                                        f"<div style='flex: 1 1 0%; min-width: 0; text-align:right; padding-right:10px; color:#475569; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; direction:rtl;'><bdi>{left_context}</bdi></div>"
                                        f"<div style='flex:none; background:#0EA5E9; color:white; padding:2px 10px; border-radius:4px; font-weight:bold; white-space:nowrap;'>{keyword}</div>"
                                        f"<div style='flex: 1 1 0%; min-width: 0; text-align:left; padding-left:10px; color:#475569; overflow:hidden; text-overflow:ellipsis; white-space:nowrap;'>{right_context}</div>"
                                        "</div>"
                                    )                                   
                                    pos_counts_html = "".join([f"<span style='background-color: {Warna_POS_Utama.get(p, '#94A3B8')}; color:white; border-radius:4px; padding:3px 8px; font-size:11px; font-weight:600; margin-right:4px;'>{p}: {c}</span>" for p, c in Counter([tok.pos_ for tok in s if tok.pos_ not in ['SPACE', 'PUNCT', 'SYM', 'X']]).items()])
                                    p_html = f"<div style='display:flex; align-items:center; flex-wrap:wrap; margin-top: 5px;'><span style='background:#10B981; color:white; padding:3px 8px; border-radius:4px; font-size:11px; font-weight:bold; margin-right:8px;'>📝 ID Segmen: {id_segmen}</span>{pos_counts_html}</div>"                                   
                                    tags_list = list(set([tok.pos_ for tok in s if tok.pos_ in deskripsi_pos]))
                                    
                                    matches_global.append({'file': fname, 'text': s.text.strip(), 'html': h_light_std, 'kwic_html': h_light_kwic, 'pills': p_html, 'tags': tags_list})
                    elif "Semantic" in mode_pencarian:
                        lang_query_wn = st.session_state.local_files[docs_terpilih[0]]['lang'] if docs_terpilih else 'en'
                        sinonim_query = dapatkan_sinonim(query_aktif, lang_query_wn)

                        for s in doc.sents:
                            if len(s.text.strip()) > 5:
                                is_sim = False
                                if any(t.lemma_.lower() in sinonim_query for t in s):
                                    is_sim = True
                                elif len(query_doc) > 0 and query_doc.has_vector:
                                    if s.has_vector and query_doc.similarity(s) >= 0.7:
                                        is_sim = True
                                    else:
                                        if any(not t.is_stop and not t.is_punct and t.has_vector and query_doc.similarity(t) >= 0.55 for t in s):
                                            is_sim = True
                                if is_sim:
                                    p_html = "<div style='display:flex; gap:5px; flex-wrap:wrap;'>" + "".join([f"<span style='background-color: {Warna_POS_Utama.get(p, '#94A3B8')}; color:white; border-radius:4px; padding:3px 8px; font-size:11px; font-weight:600;'>{p}: {c}</span>" for p, c in Counter([t.pos_ for t in s if t.pos_ not in ['SPACE', 'PUNCT', 'SYM', 'X']]).items()]) + "</div>"
                                    h_light = s.text.strip()
                                    for kata_sim in sinonim_query:
                                        h_light = re.sub(f"\\b({kata_sim})\\b", r"<mark style='background:#FDE047; padding:0 4px; border-radius:3px; color:#0F172A;'>\1</mark>", h_light, flags=re.I)
                                    matches_global.append({'file': fname, 'text': s.text.strip(), 'html': h_light, 'pills': p_html, 'tags': list(set([t.pos_ for t in s if t.pos_ in deskripsi_pos]))})
                                    
                    elif "Boolean" in mode_pencarian:
                        q_r = query_aktif.replace("AND", "&").replace("OR", "|").replace("NOT", "!")
                        for s in doc.sents:
                            lems = {t.lemma_.lower() for t in s}
                            nlp_aktif = load_ai_model(st.session_state.local_files[fname]['lang'])
                            if (" & " in q_r and all(nlp_aktif(p.strip())[0].lemma_.lower() in lems for p in q_r.split("&"))) or (" | " in q_r and any(nlp_aktif(p.strip())[0].lemma_.lower() in lems for p in q_r.split("|"))) or ("!" in q_r and nlp_aktif(q_r.replace("!", "").strip())[0].lemma_.lower() not in lems) or (query_aktif.lower() in lems):
                                h_light = s.text.strip()
                                for w in re.findall(r'\w+', query_aktif):
                                    if w.upper() not in ["AND", "OR", "NOT"]: h_light = re.sub(f"\\b({w})\\b", r"<mark style='background:#FDE047; color:#0F172A; padding:0 4px; border-radius:3px;'>\1</mark>", h_light, flags=re.I)
                                p_html = "<div style='display:flex; gap:5px; flex-wrap:wrap;'>" + "".join([f"<span style='background-color: {Warna_POS_Utama.get(p, '#94A3B8')}; color:white; border-radius:4px; padding:3px 8px; font-size:11px; font-weight:600;'>{p}: {c}</span>" for p, c in Counter([t.pos_ for t in s if t.pos_ not in ['SPACE', 'PUNCT', 'SYM', 'X']]).items()]) + "</div>"
                                matches_global.append({'file': fname, 'text': s.text.strip(), 'html': h_light, 'pills': p_html, 'tags': list(set([t.pos_ for t in s if t.pos_ in deskripsi_pos]))})

                st.session_state.search_results, st.session_state.sketch_results, st.session_state.morph_results, st.session_state.colloc_results = matches_global, {}, pd.DataFrame(), pd.DataFrame()
                st.session_state.current_page, st.session_state.last_query = 0, query_aktif

    mode_saran_aktif = ["Lemmatization", "Semantic", "Word Sketch", "Morphology", "Collocation"]
    
    query_pencarian = st.session_state.get('last_query', "").strip()   
    if query_pencarian and any(m in mode_pencarian for m in mode_saran_aktif):
        lang_saran = st.session_state.local_files[docs_terpilih[0]]['lang'] if docs_terpilih else 'en'
        query_lower = query_pencarian.lower()
        try: 
            q_lemma = nlp_query(query_pencarian)[0].lemma_.lower()
        except: 
            q_lemma = query_lower
        sinonim_set = set()
        try:
            sinonim_tambahan = dapatkan_sinonim(query_pencarian, lang_saran)
            if isinstance(sinonim_tambahan, (list, set)):
                sinonim_set.update(sinonim_tambahan)
        except Exception as e: 
            pass
        vocab_aktif = set()
        for doc_name in docs_terpilih:
            df_indeks = st.session_state.local_files[doc_name].get('index_df')
            if df_indeks is not None and not df_indeks.empty:
                vocab_aktif.update(set(df_indeks['Tokens'].explode().dropna()))
                vocab_aktif.update(set(df_indeks['Lemmas'].explode().dropna()))
            else:
                vocab_aktif.update(st.session_state.local_files[doc_name].get('vocab', []))
        saran_bersih = []
        for kata in sinonim_set:
            kata_lower = kata.lower().strip()
            
            if ' ' in kata_lower or '_' in kata_lower:
                continue
                
            if kata_lower != query_lower and kata_lower != q_lemma:
                if kata_lower in vocab_aktif:
                    saran_bersih.append(kata_lower)

        saran_kata = list(set(saran_bersih))[:8]         
        if saran_kata:
            st.markdown("<div style='font-size:14px; color:#64748B; margin-bottom:8px;'>💡 Saran kata terkait yang <b>ada di dokumen ini</b>:</div>", unsafe_allow_html=True)
            try: 
                st.pills("Saran", saran_kata, key="saran_pills_widget", on_change=aksi_klik_saran, label_visibility="collapsed")
            except AttributeError: 
                st.selectbox("💡 Saran kata terkait yang ada di dokumen:", ["-- Pilih --"] + saran_kata, key="saran_pills_widget", on_change=aksi_klik_saran)

    if "Collocation" in mode_pencarian:
        if not st.session_state.colloc_results.empty:
            df_c = st.session_state.colloc_results
            st.success(f"✅ Ditemukan **{len(df_c)}** kata sanding (collocates) untuk **'{query_aktif}'**.")
            st.dataframe(df_c, use_container_width=True, hide_index=True)
        elif query_aktif and st.session_state.last_query == query_aktif: 
            st.warning(f"🔍 Kata '{query_aktif}' tidak ditemukan atau tidak memiliki asosiasi kuat di dokumen yang dipilih.")

    elif "Word Sketch" in mode_pencarian:
        if st.session_state.sketch_results:
            st.write(f"### 🔍 Hasil Profil untuk: **{query_aktif}**")
            c_mod, c_subj, c_obj = st.columns(3)
            def r_list(header, data, color, kp):
                st.markdown(f"<div style='background:{color}; padding:10px; border-radius:8px; margin-bottom:10px;'><b>{header}</b></div>", unsafe_allow_html=True)
                if data:
                    for i, (k, jml) in enumerate(data): st.button(f"{k} ({jml}x)", key=f"{kp}_{k}_{i}", use_container_width=True, on_click=update_ws_word, args=(k,))
                else: st.caption("Tidak ditemukan.")
            with c_mod: r_list("✨ Modifiers (Sifat)", st.session_state.sketch_results["✨ Modifiers (Sifat/Penjelas)"], "#F0FDF4", "ws_mod")
            with c_subj: r_list("🏃‍♂️ Sebagai Pelaku", st.session_state.sketch_results["🏃‍♂️ Sebagai Subjek (Melakukan)"], "#EFF6FF", "ws_sub")
            with c_obj: r_list("🎯 Sebagai Korban", st.session_state.sketch_results["🎯 Sebagai Objek (Dikenai)"], "#FEF2F2", "ws_obj")
        elif query_aktif and st.session_state.last_query == query_aktif: st.warning(f"🔍 Profil kata '{query_aktif}' tidak ditemukan.")
            
    elif "Morphology" in mode_pencarian:
        if not st.session_state.morph_results.empty:
            df_g = st.session_state.morph_results
            st.success(f"✅ Ditemukan **{len(df_g)}** variasi.")
            st.altair_chart(alt.Chart(df_g).mark_bar(color="#F59E0B", cornerRadiusEnd=4).encode(x=alt.X('Frekuensi Muncul:Q'), y=alt.Y('Bentuk Teks Asli:N', sort='-x'), tooltip=['Bentuk Teks Asli', 'Kelas Kata (POS)', 'Struktur Morfologi', 'Frekuensi Muncul']).properties(height=300), use_container_width=True)
            st.dataframe(df_g, use_container_width=True, hide_index=True)
        elif query_aktif and st.session_state.last_query == query_aktif: st.warning(f"🔍 Akar kata '{query_aktif}' tidak digunakan.")
    
    else:
        if not st.session_state.search_results and query_aktif and st.session_state.last_query == query_aktif:
            st.warning(f"🔍 Pencarian '{query_aktif}' tidak ditemukan.")

        if st.session_state.search_results:
            tot_res = len(st.session_state.search_results)
            
            if "Lemmatization" in mode_pencarian:
                tampilan_lemma = st.radio("Pilih Mode Tampilan:", ["📝 Kalimat Utuh (Klasik)", "📊 KWIC (Konteks Kiri-Kanan)"], horizontal=True)
                c_i, _, c_ft, c_fd = st.columns([5, 1, 3, 1.05], gap="small")
                teks_jumlah = "baris" if "KWIC" in tampilan_lemma else "kalimat"
                c_i.markdown(f"<div style='color:#059669; font-weight:bold; padding-top:8px;'>✅ Ditemukan {tot_res} {teks_jumlah}.</div>", unsafe_allow_html=True)
            else:
                tampilan_lemma = "📝 Kalimat Utuh (Klasik)"
                c_i, _, c_ft, c_fd = st.columns([5, 1, 3, 1.05], gap="small")
                c_i.markdown(f"<div style='color:#059669; font-weight:bold; padding-top:8px;'>✅ Ditemukan {tot_res} kalimat.</div>", unsafe_allow_html=True)

            c_ft.markdown("<div style='text-align:right; padding-top:8px;'>Tampilkan per halaman:</div>", unsafe_allow_html=True)
            limit_sel = c_fd.selectbox("Limit", [str(x) for x in [5, 10, 25, 50] if x < tot_res] + [f"All ({tot_res})"], label_visibility="collapsed")
            
            IP_PAGE = tot_res if limit_sel.startswith("All") else int(limit_sel)
            tot_pages = max(1, math.ceil(tot_res / IP_PAGE))
            if st.session_state.current_page >= tot_pages: st.session_state.current_page = max(0, tot_pages - 1)
            
            s_idx = st.session_state.current_page * IP_PAGE
            
            if "Lemmatization" in mode_pencarian and "KWIC" in tampilan_lemma:
                gabungan_kwic = "<div style='width:100%; border: 1px solid #E2E8F0; border-radius:8px; background:white;'>"
                for m_data in st.session_state.search_results[s_idx:s_idx+IP_PAGE]:
                    gabungan_kwic += m_data.get('kwic_html', m_data['html']) 
                gabungan_kwic += "</div>"
                st.markdown(gabungan_kwic, unsafe_allow_html=True)
                
            else:
                for i, m_data in enumerate(st.session_state.search_results[s_idx:s_idx+IP_PAGE]):
                    with st.container(border=True):
                        html_aman = m_data['html'].replace('\n', ' ')
                        jml_kata = len(re.findall(r'\b\w+\b', m_data['text']))
                        
                        st.markdown(f"<div style='color:#334155; font-size:15.5px; margin-bottom:10px; line-height:1.6;'>{html_aman}</div>", unsafe_allow_html=True)
                        st.markdown("<hr style='margin: 5px 0 10px 0; border-color: #F1F5F9;'>", unsafe_allow_html=True)
                        
                        col_info, col_aksi = st.columns([8.5, 1.5], vertical_alignment="center")
                        
                        with col_info:
                            st.markdown(
                                f"<div style='display:flex; align-items:center; gap:8px; flex-wrap:wrap;'>"
                                f"<div style='font-size:11px; color:#0284C7; font-weight:bold; background:#F0F9FF; padding:4px 8px; border-radius:4px; border:1px solid #BAE6FD;'>📄 {m_data['file'].upper()}</div>"
                                f"<div style='font-size:11px; color:#059669; font-weight:bold; background:#ECFDF5; padding:4px 8px; border-radius:4px; border:1px solid #A7F3D0;'>🔢 {jml_kata} Kata</div>"
                                f"<div style='margin-top:-5px;'>{m_data['pills']}</div>"
                                f"</div>", 
                                unsafe_allow_html=True
                            )
                            
                        with col_aksi:
                            aksi = st.selectbox("Aksi", ["Aksi", "🏷️ POS Tag", "🌿 Syntax Tree", "🌐 Trans"], key=f"aksi_{s_idx+i}", label_visibility="collapsed")

                        if aksi == "🏷️ POS Tag":
                            st.markdown(get_colored_pos_text(load_ai_model(st.session_state.local_files[m_data['file']]['lang']), m_data['text']), unsafe_allow_html=True)
                            o_drop = [t for t in deskripsi_pos.keys() if t in m_data['tags']]
                            if o_drop: st.info(deskripsi_pos[st.selectbox("💡 Penjelasan:", o_drop, key=f"hp_{s_idx+i}")])
                        elif aksi == "🌿 Syntax Tree":
                            st.markdown(render_dependency_tree(m_data['text'], load_ai_model(st.session_state.local_files[m_data['file']]['lang'])), unsafe_allow_html=True)
                        elif aksi == "🌐 Trans":
                            c_l, c_g = st.columns([7, 3])
                            tl_name = c_l.selectbox("Ke:", list(DAFTAR_BAHASA.keys()), key=f"ln_{s_idx+i}", label_visibility="collapsed")
                            if c_g.button("Ok", key=f"g_{s_idx+i}", use_container_width=True):
                                try: st.markdown(f"<div style='background:#E0F2FE; border:1px solid #7DD3FC; padding:15px; border-radius:8px; margin-top:10px;'>{GoogleTranslator(source='auto', target=DAFTAR_BAHASA[tl_name]).translate(m_data['text'])}</div>", unsafe_allow_html=True)
                                except Exception as e: st.error(f"Error: {e}")

            st.write("") 
            c_b1, c_p, c_pi, c_n, c_b2 = st.columns([2, 1, 2, 1, 2])
            c_p.button("⬅️ Prev", use_container_width=True, disabled=(st.session_state.current_page == 0), on_click=prev_page_1)
            with c_pi:
                with st.popover(f"Hal: {st.session_state.current_page + 1} / {tot_pages}", use_container_width=True):
                    st.number_input("Pindah ke halaman:", min_value=1, max_value=max(1, tot_pages), value=st.session_state.current_page + 1, key="num_search_page", on_change=aksi_loncat_halaman_search)
            c_n.button("Next ➡️", use_container_width=True, disabled=(st.session_state.current_page >= tot_pages - 1), on_click=next_page_1)
            
# --- FRAGMENT SUMMARIZE (Hanya di tab Dokumen) ---
@st.fragment
def render_tab_summarize(docs_terpilih):
    st.markdown("<h3 style='color:#0F172A;'>📝 Analisis Dokumen Mendalam</h3>", unsafe_allow_html=True)
    target_file = st.selectbox("Pilih dokumen spesifik untuk dianalisis:", docs_terpilih, key="sel_doc_sum_sent_top")
    sub_topic, sub_sum, sub_sent = st.tabs(["🗂️ Pemodelan Topik (LDA)", "📑 Ekstraksi Ringkasan", "😊 Analisis Sentimen (VADER)"])
    
    with sub_topic:
        st.info("**💡 Fungsi:** Algoritma AI secara otomatis memindai dan mengelompokkan kata-kata ke dalam 'Tema/Topik Utama' tanpa Anda harus membaca seluruh dokumen.")
        with st.expander("📖 Panduan Detail Membaca Hasil Topik (LDA)", expanded=False):
            st.markdown("""
            * **Tujuan:** Menemukan *tema tersembunyi* dari dokumen yang sangat panjang.
            * **Cara Kerja:** Menggunakan algoritma Latent Dirichlet Allocation (LDA), AI memindai ribuan kata dan mengidentifikasi kata-kata yang sering muncul dalam konteks yang sama, lalu memasukkannya ke dalam "keranjang tema".
            * **Cara Membaca Hasil:** Setiap kotak akan menampilkan tebakan Label Topik beserta deretan kata dominan di dalamnya. Rangkai kata-kata tersebut di kepala Anda untuk menyimpulkan pembicaraannya. 
            """)
        col_t1, col_t2 = st.columns([4, 2], vertical_alignment="bottom")
        with col_t1: num_topics = st.number_input("Jumlah Topik Utama:", min_value=2, max_value=12, value=3, key="num_topics")
        with col_t2: btn_topic = st.button("🚀 Ekstrak Topik Otomatis", type="primary", use_container_width=True)

        if btn_topic:
            with st.spinner(f"Membangun model LDA tingkat lanjut untuk '{target_file}'..."):
                from sklearn.feature_extraction.text import CountVectorizer
                from sklearn.decomposition import LatentDirichletAllocation
                
                nlp_aktif = load_ai_model(st.session_state.local_files[target_file].get('lang', 'en'))
                teks_dokumen = st.session_state.local_files[target_file]['cleaned'][:30000] 

                corpus_sentences = []
                for s in nlp_aktif(teks_dokumen).sents:
                    tokens_penting = [
                        t.lemma_.lower() for t in s 
                        if t.pos_ in ['NOUN', 'PROPN', 'VERB', 'ADJ'] 
                        and not t.is_stop 
                        and not t.is_punct 
                        and len(t.text) > 2
                    ]
                    if len(tokens_penting) >= 3:
                        corpus_sentences.append(" ".join(tokens_penting))
                
                if len(corpus_sentences) < 10:
                    st.warning("⚠️ Teks bermakna terlalu pendek untuk pemodelan topik (minimal 10 kalimat).")
                else:
                    try:
                        custom_noise = {'et', 'al', 'fig', 'figure', 'table', 'use', 'using', 'based', 'study', 'result', 'analysis', 'show', 'paper'}
                        
                        vectorizer = CountVectorizer(stop_words=list(custom_noise), min_df=2, max_df=0.85)
                        dtm = vectorizer.fit_transform(corpus_sentences)
                        
                        lda = LatentDirichletAllocation(
                            n_components=num_topics, 
                            random_state=42, 
                            max_iter=25,          
                            learning_method='batch' 
                        )
                        lda.fit(dtm)
                        
                        feature_names = vectorizer.get_feature_names_out()
                        
                        st.markdown("<hr style='margin:15px 0; border:1px dashed #CBD5E1;'>", unsafe_allow_html=True)
                        cols_topik = st.columns(3) 
                        for topic_idx, topic in enumerate(lda.components_):
                            top_words = [feature_names[i] for i in topic.argsort()[:-11:-1]]
                            label_topik = f"{top_words[0].title()} & {top_words[1].title()}"
                            
                            with cols_topik[topic_idx % 3]:
                                st.markdown(f"""
                                <div style='background:#F8FAFC; border:1px solid #CBD5E1; padding:20px; border-radius:12px; margin-bottom:20px; height: 100%; box-shadow: 0 2px 4px rgba(0,0,0,0.02);'>
                                    <h4 style='color:#0369A1; margin-top:0; margin-bottom:5px; font-size:16px;'>
                                        📌 Topik {topic_idx + 1}
                                    </h4>
                                    <div style='color:#0F172A; font-weight:bold; font-size:18px; border-bottom:2px solid #BAE6FD; padding-bottom:8px; margin-bottom:10px;'>
                                        🏷️ {label_topik}
                                    </div>
                                    <div style='display:flex; flex-wrap:wrap; gap:6px;'>
                                        {''.join([f'<span style="background:white; border:1px solid #94A3B8; border-radius:12px; padding:4px 10px; font-size:13.5px; color:#334155; font-weight:500;">{w}</span>' for w in top_words])}
                                    </div>
                                </div>
                                """, unsafe_allow_html=True)
                    except Exception as e: 
                        st.error(f"Gagal mengekstrak topik: {e}")

    with sub_sum:
        st.info("**💡 Fungsi:** AI membaca teks secara keseluruhan dan menulis ulang intisarinya dengan bahasa yang lebih padat (Abstractive Summarization menggunakan BART).")
        
        with st.expander("📖 Panduan Detail Abstractive Summarization", expanded=False):
            st.markdown("""
            * **Tujuan:** Mendapatkan kesimpulan dokumen yang ditulis ulang oleh AI secara natural.
            * **Cara Kerja:** Menggunakan model **BART-Large-CNN** dari Facebook/Meta. AI tidak hanya mencomot kalimat, tetapi memahami konteks dan memparafrasekannya menjadi paragraf baru.
            * **Catatan Penting:** Model ini sangat optimal untuk **Bahasa Inggris**. Jika Anda memasukkan dokumen bahasa Indonesia, hasilnya mungkin kurang akurat atau diterjemahkan paksa.
            """)
            
        if st.button(f"🚀 Mulai Abstractive Summarize (BART)", type="primary", key="btn_run_sum"):
            with st.spinner("AI sedang membaca dan menulis ulang intisari dokumen (Proses ini mungkin memakan waktu beberapa saat)..."):
                try:
                    t_target = st.session_state.local_files[target_file]['cleaned']                      
                    tokenizer, model, device = load_bart_summarizer()
                    
                    words = t_target.split()
                    total_words = len(words)
                    
                    rasio_target = 0.15
                    
                    chunk_size = 600
                    chunks = [' '.join(words[i:i+chunk_size]) for i in range(0, total_words, chunk_size)]
                    
                    hasil_ringkasan = []
                    
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    for i, chunk in enumerate(chunks):
                        chunk_words = chunk.split()
                        jml_kata_chunk = len(chunk_words)
                        
                        if jml_kata_chunk > 50:
                            status_text.text(f"Menganalisis bagian {i+1} dari {len(chunks)}...")
                            
                            target_kata_chunk = jml_kata_chunk * rasio_target
                            max_len = max(30, int(target_kata_chunk * 1.5)) 
                            min_len = max(10, int(target_kata_chunk * 1.1))
                            
                            inputs = tokenizer(chunk, max_length=1024, return_tensors="pt", truncation=True).to(device)
                            
                            summary_ids = model.generate(
                                inputs["input_ids"], 
                                max_length=max_len, 
                                min_length=min_len, 
                                do_sample=False,
                                num_beams=1 
                            )
                           
                            summary_text = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
                            hasil_ringkasan.append(summary_text)
                        
                        progress_bar.progress((i + 1) / len(chunks))
                    
                    status_text.empty()
                    progress_bar.empty()
                    
                    st.session_state.summary_results[target_file] = "\n\n".join(hasil_ringkasan)
                    
                    st.success("✅ Penulisan ulang kesimpulan selesai!")
                except Exception as e: 
                    st.error(f"Error saat merangkum: {e}")

        if target_file in st.session_state.summary_results:
            t_hasil = st.session_state.summary_results[target_file]
            kalimat_list = [k for k in t_hasil.replace('?', '.').replace('!', '.').split('.') if len(k.strip()) > 2]
            jml_kalimat = len(kalimat_list)

            with st.container(border=True):
                cj, ca = st.columns([0.9, 0.1], vertical_alignment="center")
                with cj: 
                    st.markdown(f"#### 📑 Teks Ringkasan (BART AI)")
                with ca:
                    with st.popover("⚙️", use_container_width=True):
                        st.markdown("<p style='font-size: 14px; font-weight: bold; margin-bottom: 5px;'>Unduh Teks:</p>", unsafe_allow_html=True)
                        st.download_button("📄 Format TXT", data=t_hasil.encode('utf-8'), file_name=f"sum_bart_{target_file}.txt", use_container_width=True)
                        try:
                            doc_ex = Document(); doc_ex.add_heading(f"Ringkasan BART", 0)
                            for p in t_hasil.split('\n\n'):
                                if p.strip(): doc_ex.add_paragraph(p.strip())
                            b_docx = io.BytesIO(); doc_ex.save(b_docx)
                            st.download_button("📝 Format DOCX", data=b_docx.getvalue(), file_name=f"sum_bart_{target_file}.docx", use_container_width=True)
                        except: 
                            pass
                            
                st.markdown("<hr style='margin: 10px 0; opacity: 0.2;'>", unsafe_allow_html=True)                   
                st.markdown(f"<div style='font-size: 13px; color: #64748B; margin-bottom: 15px;'>📊 <b>Statistik:</b> AI menghasilkan {jml_kalimat} kalimat ringkasan baru.</div>", unsafe_allow_html=True)                   
                html_paragraphs = "".join([f"<p style='margin-bottom: 15px; text-indent: 30px;'>{p.strip()}</p>" for p in t_hasil.split('\n\n') if p.strip()])
                st.markdown(f"<div style='color: #334155; font-size: 16px; line-height: 1.8; text-align: justify;'>{html_paragraphs}</div>", unsafe_allow_html=True)
            
            st.write("")
            ak_sum = st.selectbox("Analisis Lanjutan:", ["Pilih Aksi...", "🏷️ POS Tagging", "🌐 Translate"], key="ak_s_l")
            
            if ak_sum == "🏷️ POS Tagging":
                with st.spinner("Membedah Tagging dan Menghitung..."): 
                    nlp_model = load_ai_model(st.session_state.local_files[target_file].get('lang', 'en'))
                    doc_pos = nlp_model(t_hasil)
                    pos_counts = {}
                    for token in doc_pos:
                        if token.pos_ not in ['PUNCT', 'SPACE']: 
                            pos_counts[token.pos_] = pos_counts.get(token.pos_, 0) + 1
                    badges_html = " ".join([f"<span style='background-color: #E2E8F0; color: #334155; padding: 4px 10px; border-radius: 8px; font-size: 13px; margin-right: 6px; margin-bottom: 8px; display: inline-block;'><b>{pos}</b>: {count}</span>" for pos, count in sorted(pos_counts.items(), key=lambda x: x[1], reverse=True)])
                    
                    st.markdown(f"<div style='margin-bottom: 15px;'>{badges_html}</div>", unsafe_allow_html=True)
                    pos_html = get_colored_pos_text(nlp_model, t_hasil)
                    st.markdown(f"<div style='margin-top: 5px; padding: 10px; display: block; overflow: hidden; background-color: transparent;'>{pos_html}</div>", unsafe_allow_html=True)                
            elif ak_sum == "🌐 Translate":
                cls, cgs, _ = st.columns([2, 1, 5])
                with cls: 
                    ts_code = DAFTAR_BAHASA[st.selectbox("Ke:", list(DAFTAR_BAHASA.keys()), key="l_s_u", label_visibility="collapsed")]
                with cgs: 
                    if st.button("Ok", key="g_s_u", use_container_width=True):
                        with st.spinner("Translating..."):
                            try: 
                                translated_text = GoogleTranslator(source='auto', target=ts_code).translate(t_hasil[:4500])
                                trans_paragraphs = "".join([f"<p style='margin-bottom: 15px;'>{p.strip()}</p>" for p in translated_text.split('\n') if p.strip()])
                                st.markdown(f"<div style='background-color:#F0FDF4; padding:25px; border-radius:12px; border: 1px solid #BBF7D0; margin-top:15px; color: #166534; font-size: 16px; line-height: 1.8; text-align: justify;'>{trans_paragraphs}</div>", unsafe_allow_html=True)
                            except Exception as e: 
                                st.error(f"Gagal terjemahkan: {e}")

    with sub_sent:
        st.info("**💡 Fungsi:** AI mengukur nada emosi dan polaritas (Positif, Negatif, Netral) dari setiap kalimat di dalam dokumen.")
        
        with st.expander("📖 Panduan Detail Analisis Sentimen (VADER)", expanded=False):
            st.markdown("""
            * **Tujuan:** Mengetahui kecenderungan emosi penulis atau subjek yang sedang dibahas di dalam teks secara statistik.
            * **Cara Kerja:** Menggunakan kamus sentimen **VADER**, AI memberikan bobot emosi pada tiap kata. Jika dokumen bukan berbahasa Inggris, teks akan diterjemahkan sementara di latar belakang agar deteksi emosinya tetap akurat.
            * **Cara Membaca Hasil:** * **Skor Compound:** Metrik gabungan yang bernilai dari -1 (Sangat Negatif) hingga +1 (Sangat Positif). Kalimat dianggap positif jika skor ≥ 0.05.
                * **Eksplorasi Kalimat:** Anda bisa mengklik/menceklis baris di dalam tabel untuk memfilter dan membaca langsung kalimat mana saja yang bersentimen negatif atau positif.
            """)
            
        if 'sentiment_results' not in st.session_state:
            st.session_state.sentiment_results = {}
        if st.button(f"🎭 Mulai Analisis Sentimen", type="primary", key="btn_run_sent"):
            with st.spinner(f"Menganalisis emosi 100% kalimat di '{target_file}'..."):
                sia = SentimentIntensityAnalyzer()
                kalimat_semua = nltk.sent_tokenize(st.session_state.local_files[target_file]['cleaned'])
                
                data_sentimen = []
                lang_doc_sent = st.session_state.local_files[target_file].get('lang', 'en')
                
                for teks_asli in kalimat_semua:
                    teks_asli = teks_asli.strip()
                    if len(teks_asli) > 15:
                        teks_vader = teks_asli
                        if lang_doc_sent != 'en':
                            try: teks_vader = GoogleTranslator(source='auto', target='en').translate(teks_asli)
                            except: pass
                        skor = sia.polarity_scores(teks_vader)
                        data_sentimen.append({"Kalimat": teks_asli, "Positif": skor['pos'], "Negatif": skor['neg'], "Netral": skor['neu'], "Compound": skor['compound']})
                
                if data_sentimen:
                    df_sent = pd.DataFrame(data_sentimen)
                    df_sent['Kategori'] = df_sent['Compound'].apply(lambda c: 'Positif' if c >= 0.05 else ('Negatif' if c <= -0.05 else 'Netral'))
                    st.session_state.sentiment_results[target_file] = df_sent

        if target_file in st.session_state.sentiment_results:
            df_sent = st.session_state.sentiment_results[target_file]
            
            avg_compound = df_sent['Compound'].mean()
            status, bg_color, txt_color = ("Positif 😊", "#DCFCE7", "#166534") if avg_compound >= 0.05 else (("Negatif 😠", "#FEE2E2", "#991B1B") if avg_compound <= -0.05 else ("Netral 😐", "#F1F5F9", "#334155"))
            
            st.markdown(f"<div style='background-color: {bg_color}; padding: 20px; border-radius: 12px; text-align: center; margin-bottom: 20px; border: 1px solid {txt_color}44;'><h4 style='color: {txt_color}; margin: 0;'>Sentimen Keseluruhan</h4><h1 style='color: {txt_color}; margin: 10px 0 0 0; font-size: 42px;'>{status}</h1><p style='color: {txt_color}; margin: 5px 0 0 0;'>Skor: <b>{avg_compound:.3f}</b></p></div>", unsafe_allow_html=True)
            
            distribusi = df_sent['Kategori'].value_counts().reset_index()
            distribusi.columns = ['Kategori', 'Jumlah Kalimat']
            total_kalimat = distribusi['Jumlah Kalimat'].sum()
            distribusi['Persentase'] = (distribusi['Jumlah Kalimat'] / total_kalimat * 100).apply(lambda x: f"{x:.1f}%")
            
            col_pie, col_ket = st.columns([1, 1])
            with col_pie: 
                base = alt.Chart(distribusi).encode(
                    theta=alt.Theta(field="Jumlah Kalimat", type="quantitative"),
                    color=alt.Color(field="Kategori", type="nominal", scale=alt.Scale(domain=['Positif', 'Netral', 'Negatif'], range=['#10B981', '#94A3B8', '#EF4444']))
                )
                pie = base.mark_arc(innerRadius=50).encode(tooltip=[alt.Tooltip('Kategori:N'), alt.Tooltip('Jumlah Kalimat:Q'), alt.Tooltip('Persentase:N')])
                st.altair_chart((pie).properties(height=300), use_container_width=True)
                
            with col_ket: 
                st.markdown("<div style='font-size:13px; color:#64748B; margin-bottom:5px;'>👆 <b>Interaktif:</b> Ceklis satu atau beberapa baris tabel di bawah untuk mengeksplorasi kalimat.</div>", unsafe_allow_html=True)     
                tabel_event = st.dataframe(
                    distribusi, 
                    use_container_width=True, 
                    hide_index=True,
                    on_select="rerun", 
                    selection_mode="multi-row" 
                )

            st.markdown("<hr style='border: 1px dashed #E2E8F0; margin: 20px 0;'>", unsafe_allow_html=True)
            
            def aksi_prev_sent(state_key): st.session_state[state_key] -= 1
            def aksi_next_sent(state_key): st.session_state[state_key] += 1
            def aksi_loncat_sent(state_key, input_key): st.session_state[state_key] = st.session_state[input_key] - 1

            if len(tabel_event.selection.rows) > 0:
                idx_terpilih = tabel_event.selection.rows
                kategori_terpilih_list = distribusi.iloc[idx_terpilih]['Kategori'].tolist()
                
                st.markdown(f"<h4 style='color:#0F172A;'>📋 Eksplorasi Kalimat Terpilih</h4>", unsafe_allow_html=True)
                kolom_eksplorasi = st.columns(len(kategori_terpilih_list))
                
                for i, kat_aktif in enumerate(kategori_terpilih_list):
                    with kolom_eksplorasi[i]:
                        p_key = f"sent_p_{target_file}_{kat_aktif}"
                        
                        if p_key not in st.session_state:
                            st.session_state[p_key] = 0
                            
                        p = st.session_state[p_key]
                        
                        df_filter = df_sent[df_sent['Kategori'] == kat_aktif].reset_index(drop=True)
                        tot = len(df_filter)
                        
                        if tot > 0:
                            if p >= tot: p = tot - 1
                            if p < 0: p = 0
                            
                            row = df_filter.iloc[p]
                            
                            warna_garis = '#10B981' if kat_aktif == 'Positif' else ('#EF4444' if kat_aktif == 'Negatif' else '#94A3B8')
                            bg_warna = '#ECFDF5' if kat_aktif == 'Positif' else ('#FEF2F2' if kat_aktif == 'Negatif' else '#F8FAFC')
                            teks_warna = '#064E3B' if kat_aktif == 'Positif' else ('#7F1D1D' if kat_aktif == 'Negatif' else '#334155')
                            
                            st.markdown(
                                f"<div style='background-color: {bg_warna}; border: 1px solid {warna_garis}; border-radius: 8px; padding: 15px; margin-bottom: 10px; height: 100%; min-height: 140px;'>"
                                f"<h6 style='color:{warna_garis}; margin:0 0 8px 0; font-size:13px; border-bottom: 1px solid {warna_garis}44; padding-bottom:5px;'>{kat_aktif} | Skor: {row['Compound']:.3f}</h6>"
                                f"<div style='color: {teks_warna}; font-size: 14.5px; line-height: 1.6;'>\"{row['Kalimat']}\"</div>"
                                f"</div>", unsafe_allow_html=True
                            )
                            
                            c_p, c_i, c_n = st.columns([2, 3, 2])
                            c_p.button("⬅️", disabled=(p == 0), key=f"prev_s_{target_file}_{kat_aktif}", use_container_width=True, on_click=aksi_prev_sent, args=(p_key,))
                            
                            with c_i:
                                with st.popover(f"{p+1}/{tot}", use_container_width=True):
                                    st.number_input("Hal:", min_value=1, max_value=max(1, tot), value=p+1, key=f"num_{p_key}", on_change=aksi_loncat_sent, args=(p_key, f"num_{p_key}"))
                                    
                            c_n.button("➡️", disabled=(p == tot - 1), key=f"next_s_{target_file}_{kat_aktif}", use_container_width=True, on_click=aksi_next_sent, args=(p_key,))
                
                st.markdown("<hr style='border: 1px dashed #CBD5E1; margin: 30px 0;'>", unsafe_allow_html=True)
            
            col_pos, col_neg = st.columns(2)
            with col_pos:
                st.markdown("<h4 style='color:#10B981;'>📈 5 Kalimat Positif Teratas</h4>", unsafe_allow_html=True)
                for _, row in df_sent.nlargest(5, 'Compound').iterrows(): st.markdown(f"<div style='background:#ECFDF5; padding:10px; border-radius:6px; border-left:4px solid #10B981; margin-bottom:8px; font-size:14px; color:#064E3B;'>{row['Kalimat']} <br><small><b>Skor: +{row['Compound']:.2f}</b></small></div>", unsafe_allow_html=True)
            with col_neg:
                st.markdown("<h4 style='color:#EF4444;'>📉 5 Kalimat Negatif Teratas</h4>", unsafe_allow_html=True)
                for _, row in df_sent.nsmallest(5, 'Compound').iterrows(): st.markdown(f"<div style='background:#FEF2F2; padding:10px; border-radius:6px; border-left:4px solid #EF4444; margin-bottom:8px; font-size:14px; color:#7F1D1D;'>{row['Kalimat']} <br><small><b>Skor: {row['Compound']:.2f}</b></small></div>", unsafe_allow_html=True)
        else: 
            st.warning("Tidak ada kalimat yang bisa dianalisis.")

# --- FRAGMENT INTERLINEAR GLOSS ---
@st.fragment
def render_tab_gloss(docs_terpilih):
    st.markdown("<h3 style='color:#0F172A;'>📖 Interlinear Glossing</h3>", unsafe_allow_html=True)
    st.info("**💡 Fungsi:** Membedah struktur kalimat secara mendetail dengan menyejajarkan teks asli, kata dasar, kelas kata, dan fitur morfologinya secara vertikal per kata.")
    
    with st.expander("📖 Panduan Detail Interlinear Gloss", expanded=False):
        st.markdown("""
        * **Tujuan:** Analisis linguistik formal dan pembelajaran struktur bahasa.
        * **Cara Kerja:** Sistem memecah kalimat menjadi token (kata), lalu AI memberikan anotasi untuk masing-masing token. Teks disusun secara menurun dari atas ke bawah untuk setiap kata.
        * **Membaca Hasil:** 1. **Teks Asli** (Paling atas)
            2. **Lemma** (Kata dasar)
            3. **POS Tag** (Kelas kata dengan warna khusus)
            4. **Morfologi** (Bentuk grammar, seperti *Tense*, *Gender*, *Number*, dll)
        * Di bagian paling bawah akan tersedia **Free Translation** (Terjemahan Bebas) dari kalimat utuh tersebut.
        """)

    col_g1, col_g2 = st.columns([6, 4])
    with col_g1: target_file = st.selectbox("📄 Pilih dokumen:", docs_terpilih, key="sel_doc_gloss")
    with col_g2: target_lang_trans = st.selectbox("🌐 Terjemahan Bebas (Free Translation) ke:", list(DAFTAR_BAHASA.keys()), index=0, key="sel_lang_gloss")

    if st.button("🚀 Buat Interlinear Gloss", type="primary", key="btn_run_gloss", use_container_width=True):
        st.session_state.run_gloss = True
        st.session_state.gloss_page = 0
        
    if st.session_state.get('run_gloss', False):
        with st.spinner("Menganalisis dan menyejajarkan struktur seluruh kalimat..."):
            file_data = st.session_state.local_files[target_file]
            nlp_model = load_ai_model(file_data.get('lang', 'en'))
            
            doc = nlp_model(file_data['cleaned'])
            sentences = [s for s in doc.sents if len(s.text.strip()) > 5]
            
            st.session_state.gloss_sents = sentences
            
    if 'gloss_sents' in st.session_state and st.session_state.gloss_sents:
        sents = st.session_state.gloss_sents
        tot = len(sents)
        
        p_key = "gloss_page"
        if p_key not in st.session_state: st.session_state[p_key] = 0
        p = st.session_state[p_key]
        
        if p >= tot: p = tot - 1
        if p < 0: p = 0
        
        s = sents[p]
        
        gloss_html = "<div style='display: flex; flex-wrap: wrap; gap: 20px; margin-bottom: 20px; background: white; padding: 25px; border-radius: 12px; border: 1px solid #E2E8F0; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);'>"
        
        for token in s:
            if token.is_space or token.is_punct: continue
            
            morph_str = str(token.morph).replace("|", "<br>") if str(token.morph) else "-"
            bg_warna = Warna_POS_Utama.get(token.pos_, "#94A3B8")
            
            gloss_html += "<div style='display: flex; flex-direction: column; align-items: flex-start; min-width: 70px;'>"
            gloss_html += f"<span style='font-size: 17px; font-weight: 800; color: #0F172A; margin-bottom: 4px;'>{token.text}</span>"
            gloss_html += f"<span style='font-size: 14px; color: #0284C7; margin-bottom: 6px;'><i>{token.lemma_}</i></span>"
            gloss_html += f"<span style='font-size: 11px; font-weight: bold; color: white; background: {bg_warna}; padding: 3px 8px; border-radius: 4px; margin-bottom: 6px;'>{token.pos_}</span>"
            gloss_html += f"<span style='font-size: 11px; color: #64748B; line-height: 1.3;'>{morph_str}</span>"
            gloss_html += "</div>"
            
        gloss_html += "</div>"
        
        try:
            free_translation = GoogleTranslator(source='auto', target=DAFTAR_BAHASA[target_lang_trans]).translate(s.text.strip())
        except:
            free_translation = "Gagal menerjemahkan (Cek koneksi internet Anda)."
            
        st.markdown(f"<h5 style='color:#334155; margin-top:20px;'>Kalimat {p+1} / {tot}</h5>", unsafe_allow_html=True)
        st.markdown(gloss_html, unsafe_allow_html=True)
        st.markdown(f"""
        <div style='background: #F8FAFC; padding: 20px; border-radius: 12px; border-left: 5px solid #0EA5E9; border-top: 1px solid #E2E8F0; border-right: 1px solid #E2E8F0; border-bottom: 1px solid #E2E8F0;'>
            <p style='margin:0 0 5px 0; font-size:13px; color:#64748B; font-weight:600; text-transform:uppercase;'>Free Translation:</p>
            <p style='margin:0; font-size:16px; color:#0F172A; line-height:1.6;'>{free_translation}</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.write("")
        c_p, c_i, c_n = st.columns([2, 3, 2])
        
        def prev_g(): st.session_state[p_key] -= 1
        def next_g(): st.session_state[p_key] += 1
        def jump_g(): st.session_state[p_key] = st.session_state["num_jump_g"] - 1
        
        c_p.button("⬅️ Sebelumnya", disabled=(p == 0), key="prev_gloss", use_container_width=True, on_click=prev_g)
        with c_i:
            with st.popover(f"Lompat ke Hal: {p+1}/{tot}", use_container_width=True):
                st.number_input("Kalimat ke:", min_value=1, max_value=max(1, tot), value=p+1, key="num_jump_g", on_change=jump_g)
        c_n.button("Selanjutnya ➡️", disabled=(p == tot - 1), key="next_gloss", use_container_width=True, on_click=next_g)

@st.fragment
def render_tab_transkrip(docs_terpilih):
    st.markdown("<h3 style='color:#0F172A;'>⏱️ Timeline Transkrip Audio</h3>", unsafe_allow_html=True)
    ada_voice = False
    
    for fname in docs_terpilih:
        file_data = st.session_state.local_files.get(fname)
        if file_data and file_data.get('type') == 'voice' and file_data.get('segments'):
            ada_voice = True
            
            full_transcript_formatted = []
            current_speaker = 1
            
            for idx, seg in enumerate(file_data['segments']):
                if idx > 0 and (seg['start'] - file_data['segments'][idx-1]['end'] > 2.0):
                    current_speaker = 2 if current_speaker == 1 else 1
                
                label = f"Orang {current_speaker}"
                timestamp = f"[{format_detik_ke_jam(seg['start'])} - {format_detik_ke_jam(seg['end'])}]"
                line = f"{label} {timestamp}: {seg['text'].strip()}"
                full_transcript_formatted.append(line)

            text_untuk_download = "\n".join(full_transcript_formatted)

            col_h1, col_h2 = st.columns([7, 3])
            with col_h1:
                st.markdown(f"**Berkas:** `{fname}`")
            with col_h2:
                st.download_button(
                    label="📥 Unduh Transkrip Utuh",
                    data=text_untuk_download,
                    file_name=f"Full_Transkrip_{fname.split('.')[0]}.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key=f"dl_full_{fname}"
                )

            with st.expander(f"🎙️ Lihat Detail Percakapan: {fname}", expanded=True):
                current_speaker = 1
                
                for idx, seg in enumerate(file_data['segments'], start=1):
                    if idx > 1 and (seg['start'] - file_data['segments'][idx-2]['end'] > 2.0):
                        current_speaker = 2 if current_speaker == 1 else 1
                    
                    mulai = format_detik_ke_jam(seg['start'])
                    selesai = format_detik_ke_jam(seg['end'])
                    teks_segmen = seg['text'].strip()
                    warna_label = "#0EA5E9" if current_speaker == 1 else "#F59E0B"
                    
                    st.markdown(f"""
                    <div style='margin-bottom: 0px; padding: 12px 15px; background-color: #FFFFFF; border-left: 5px solid {warna_label}; border-radius: 6px 6px 0 0; border: 1px solid #E2E8F0; border-bottom: none;'>
                        <div style='display: flex; justify-content: space-between; margin-bottom: 5px;'>
                            <span style='color: {warna_label}; font-weight: 800; font-size: 13px;'>👤 ORANG {current_speaker}</span>
                            <span style='color: #64748b; font-size: 12px; font-weight: 600;'>{mulai} - {selesai}</span>
                        </div>
                        <span style='color: #1E293B; font-size: 15px; line-height: 1.5;'>"{teks_segmen}"</span>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if file_data.get('audio_bytes'):
                        try:
                            from pydub import AudioSegment
                            import io
                            
                            audio_io = io.BytesIO(file_data['audio_bytes'])
                            audio_file = AudioSegment.from_file(audio_io)
                            potongan = audio_file[int(seg['start'] * 1000) : int(seg['end'] * 1000)]
                            
                            out_io = io.BytesIO()
                            potongan.export(out_io, format="mp3")
                            audio_bytes_slice = out_io.getvalue()
                            
                            st.markdown("<div style='background-color:#F8FAFC; padding:10px 15px; border-radius: 0 0 6px 6px; border: 1px solid #E2E8F0; margin-bottom:15px;'>", unsafe_allow_html=True)
                            c_player, c_dl_a, c_dl_t = st.columns([5, 2.5, 2.5], vertical_alignment="center")
                            
                            with c_player:
                                st.audio(audio_bytes_slice, format="audio/mp3")
                            with c_dl_a:
                                st.download_button(label="🎵 MP3", data=audio_bytes_slice, 
                                                 file_name=f"Orang{current_speaker}_Part{idx}.mp3", 
                                                 mime="audio/mp3", use_container_width=True, key=f"dla_{fname}_{idx}")
                            with c_dl_t:
                                teks_dl = f"Orang {current_speaker} [{mulai}]: {teks_segmen}"
                                st.download_button(label="📄 Teks", data=teks_dl.encode('utf-8'), 
                                                 file_name=f"Orang{current_speaker}_Part{idx}.txt", 
                                                 mime="text/plain", use_container_width=True, key=f"dlt_{fname}_{idx}")
                            st.markdown("</div>", unsafe_allow_html=True)
                        except:
                            st.warning("Gagal memproses potongan audio.")
    if not ada_voice:
        st.info("ℹ️ Pilih file audio di Sidebar untuk melihat transkrip berformat narasumber.")

# ==========================================
# 6. RENDER DOKUMEN WORKSPACE
# ==========================================
@st.fragment
def render_workspace_dokumen(file_names_doc):
    st.markdown("<div style='font-size:14px; font-weight:600; color:#475569; margin-bottom:8px;'>Pilih & Filter Dokumen Aktif:</div>", unsafe_allow_html=True) 
    if 'ms_doc' not in st.session_state:
        st.session_state.ms_doc = file_names_doc[:1] if file_names_doc else []

    col_f1, col_s1, col_s2 = st.columns([3, 5, 2], vertical_alignment="bottom")
    
    with col_f1:
        pilihan_grup = st.selectbox("Berdasarkan Sub-Corpus:", ["Semua File"] + list(st.session_state.sub_corpora.keys()), label_visibility="collapsed", key="pg_doc")
    with col_s2:
        cek_semua = st.checkbox("Pilih Semua", key="ca_doc")
    with col_s1:
        tersedia = file_names_doc if pilihan_grup == "Semua File" else [f for f in st.session_state.sub_corpora.get(pilihan_grup, []) if f in file_names_doc]
        if cek_semua:
            st.session_state.ms_doc = tersedia

        with st.popover(f"📂 Pilih File ({len([f for f in st.session_state.ms_doc if f in tersedia])} Terpilih)", use_container_width=True):
            temp_selection = []
            for f in tersedia:
                if st.checkbox(f, value=(f in st.session_state.ms_doc), key=f"chk_doc_{f}"):
                    temp_selection.append(f)
            
            if not cek_semua:
                st.session_state.ms_doc = temp_selection
                
    selected_docs = [doc for doc in st.session_state.ms_doc if doc in tersedia]

    if selected_docs:
        st.markdown("---")
        tab_c_doc, tab_s_doc, tab_sum_doc, tab_gloss_doc = st.tabs(["⚖️ Perbandingan dan Visual", "🔍 Pencarian", "📝 Rangkuman", "📖 Interlinear Gloss"])
        
        with tab_c_doc: 
            render_tab_compare(selected_docs, suffix="doc")
            st.markdown("<hr style='border: 2px dashed #CBD5E1; margin: 40px 0;'>", unsafe_allow_html=True)
            render_tab_visual(selected_docs, suffix="doc")
            
        with tab_s_doc: render_tab_search(selected_docs)
        with tab_sum_doc: render_tab_summarize(selected_docs)
        
        with tab_gloss_doc: render_tab_gloss(selected_docs)
    else:
        st.warning("⚠️ Silakan pilih minimal 1 dokumen untuk dianalisis.")


with tab_induk_doc:
    uploaded_docs = st.file_uploader(
        "Upload Dokumen Teks & ELAN (Support: PDF, DOCX, TXT, EAF)", 
        accept_multiple_files=True, 
        type=['pdf', 'docx', 'txt', 'eaf'],
        key=f"uploader_doc_{st.session_state.get('uploader_key', 'default')}"
    )
    if uploaded_docs:
        for doc in uploaded_docs:
            if doc.name.endswith('.eaf'):
                df_dummy = pd.DataFrame({'Kolom1': [1, 2], 'Kolom2': ['A', 'B']}) 
                nama_key = f"df_{doc.name}"
                st.session_state.local_files[nama_key] = df_dummy

    eaf_dfs = {k: v for k, v in st.session_state.local_files.items() if k.startswith("df_") and k.endswith(".eaf")}

    
    if eaf_dfs:
        st.markdown("<div style='background:#F0FDF4; border:1px solid #BBF7D0; padding:15px; border-radius:8px; margin-bottom:20px;'>", unsafe_allow_html=True)
        st.markdown("<h4 style='margin-top:0; color:#166534;'>📥 Download Hasil Konversi ELAN</h4>", unsafe_allow_html=True)
        for eaf_key, df_eaf in eaf_dfs.items():
            nama_file_asli = eaf_key.replace("df_", "")
            buffer_excel = io.BytesIO()
            with pd.ExcelWriter(buffer_excel, engine='openpyxl') as writer:
                df_eaf.to_excel(writer, index=False, sheet_name='Data_EAF')
            
            st.download_button(
                label=f"📊 Unduh Excel: {nama_file_asli}", 
                data=buffer_excel.getvalue(), 
                file_name=f"konversi_{nama_file_asli.replace('.eaf', '')}.xlsx", 
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                key=f"dl_eaf_{nama_file_asli}"
            )
        st.markdown("</div>", unsafe_allow_html=True)
        
    if uploaded_docs:
        file_baru_diproses = [f for f in uploaded_docs if f.name not in st.session_state.local_files]
        if file_baru_diproses:
            ada_file_baru = True
            total_file = len(file_baru_diproses)
            
            bar_progres = st.progress(0, text="Memulai pemrosesan dokumen...")
            
            for i, file in enumerate(file_baru_diproses):
                base_pct = (i / total_file) * 100
                def hitung_pct(step_pct): return int(base_pct + (step_pct / total_file))
                file_extension = file.name.split('.')[-1].lower()
                bar_progres.progress(hitung_pct(10), text=f"📄 [1/4] Membaca teks dari '{file.name}'...")
                raw_text = ""
                df_eaf_temp = None 
                if file_extension == 'eaf':
                    eaf_content = file.getvalue().decode('utf-8')
                    df_eaf_temp = process_eaf_ultra_clean(eaf_content)
                    if df_eaf_temp is not None:
                        raw_text = "\n".join(df_eaf_temp['Source_Sentence'].dropna().astype(str).tolist())
                        st.session_state.local_files[f"df_{file.name}"] = df_eaf_temp
                else:
                    raw_text = extract_text(file)
                
                bar_progres.progress(hitung_pct(40), text=f"🧹 [2/4] Membersihkan format teks '{file.name}'...")
                teks_bersih = bersihkan_teks_untuk_analisis(raw_text)
                
                bar_progres.progress(hitung_pct(60), text=f"🌐 [3/4] Mendeteksi bahasa '{file.name}'...")
                try: deteksi_lang = detect(teks_bersih[:5000]) 
                except: deteksi_lang = 'en'

                bar_progres.progress(hitung_pct(70), text=f"🗂️ Membangun Indeks Cache untuk '{file.name}'...")
                nlp_aktif = load_ai_model(deteksi_lang)
                df_indeks = pd.DataFrame()
                
                if file_extension == 'eaf' and df_eaf_temp is not None:
                    df_indeks = buat_indeks_dokumen(file.name, df_eaf_temp, nlp_aktif, 'eaf')
                else:
                    df_indeks = buat_indeks_dokumen(file.name, teks_bersih, nlp_aktif, 'txt')

                bar_progres.progress(hitung_pct(80), text=f"🧩 [4/4] Mengekstrak kosakata & statistik '{file.name}'...")
                
                pola_kata = r'\b[a-zA-Z0-9]+(?:-[a-zA-Z0-9]+)*\b'
                semua_kata = re.findall(pola_kata, teks_bersih.lower())
                
                try:
                    jml_kalimat = len(nltk.sent_tokenize(teks_bersih))
                except:
                    jml_kalimat = teks_bersih.count('.') + teks_bersih.count('!') + teks_bersih.count('?')
                
                st.session_state.local_files[file.name] = {
                    'text': raw_text, 
                    'cleaned': teks_bersih, 
                    'lang': deteksi_lang,
                    'vocab': set(semua_kata), 
                    'type': 'doc', 
                    'stats': {'k': jml_kalimat, 'w': len(semua_kata)},
                    'index_df': df_indeks
                }                
                bar_progres.progress(hitung_pct(100), text=f"✅ Selesai memproses '{file.name}'!")
                
            bar_progres.empty()
            st.session_state.uploader_key += 1
            st.rerun()

    doc_files = {k: v for k, v in st.session_state.local_files.items() if v.get('type') == 'doc'}
    
    if doc_files:
        file_names_doc = list(doc_files.keys())
        with st.expander("📂 Manajemen Korpus & Sub-Corpora (Dokumen)", expanded=False):
            t_manage, t_group = st.tabs(["📄 Daftar File", "🗂️ Kelola Sub-Corpora"])
            with t_manage:
                corpus_data_doc = []
                for f_name, f_data in doc_files.items():
                    grup_file = [g for g, files in st.session_state.sub_corpora.items() if f_name in files]
                    corpus_data_doc.append({
                        "Hapus": False, 
                        "Nama File": f_name, 
                        "Bahasa": f_data.get('lang', 'en').upper(),
                        "Sub-Corpus": ", ".join(grup_file) if grup_file else "Unassigned",
                        "Total Kalimat": f_data['stats']['k'],  
                        "Total Kata": f_data['stats']['w'],
                        "Kekayaan Kata (%)": round((len(f_data['vocab']) / f_data['stats']['w']) * 100, 2) if f_data['stats']['w'] > 0 else 0,
                    })
                df_corpus_doc = pd.DataFrame(corpus_data_doc)
                
                edited_df_doc = st.data_editor(
                    df_corpus_doc, 
                    column_config={
                        "Hapus": st.column_config.CheckboxColumn("❌ Hapus", default=False), 
                        "Nama File": st.column_config.TextColumn("📄 Dokumen", disabled=True, width="medium"), 
                        "Bahasa": st.column_config.TextColumn("🌐 BHS", disabled=True), 
                        "Sub-Corpus": st.column_config.TextColumn("🗂️ Group", disabled=True), 
                        "Total Kalimat": st.column_config.NumberColumn("📝 Total Kalimat", disabled=True),
                        "Total Kata": st.column_config.NumberColumn("📊 Total Kata", disabled=True),
                        "Kekayaan Kata (%)": st.column_config.ProgressColumn("💎 Richness", format="%.2f%%", min_value=0, max_value=100)
                    }, 
                    hide_index=True, 
                    use_container_width=True, 
                    key="ed_doc"
                )
                
                file_untuk_dihapus = edited_df_doc[edited_df_doc["Hapus"] == True]["Nama File"].tolist()
                if st.button("🗑️ Hapus File Terpilih", type="primary", disabled=not file_untuk_dihapus, key="del_doc"):
                    for fname in file_untuk_dihapus:
                        del st.session_state.local_files[fname]
                        for g in st.session_state.sub_corpora:
                            if fname in st.session_state.sub_corpora[g]: st.session_state.sub_corpora[g].remove(fname)
                    st.rerun()
            with t_group:
                c1, c2 = st.columns([1, 1])
                with c1:
                    new_group = st.text_input("Buat Sub-Corpus Baru:", placeholder="Contoh: Paper_NLP_2026", key="ng_doc")
                    if st.button("➕ Tambah Grup", key="addg_doc") and new_group:
                        if new_group not in st.session_state.sub_corpora:
                            st.session_state.sub_corpora[new_group] = []
                            st.success(f"Grup '{new_group}' dibuat!")
                with c2:
                    target_grup = st.selectbox("Pilih Grup Tujuan:", list(st.session_state.sub_corpora.keys()), key="tg_doc")
                    files_to_add = st.multiselect("Pilih File untuk Dimasukkan:", file_names_doc, key="ms_add_doc")
                    if st.button("📥 Masukkan ke Grup", key="btn_add_doc"):
                        for f in files_to_add:
                            if f not in st.session_state.sub_corpora[target_grup]: st.session_state.sub_corpora[target_grup].append(f)
                        st.toast(f"Berhasil update grup {target_grup}!")
                        st.rerun()

        render_workspace_dokumen(file_names_doc)
    else:
        st.info("👋 Silakan upload dokumen teks untuk mulai menggunakan fitur Dokumen.")

# ==========================================
# 7. RENDER VOICE WORKSPACE
# ==========================================

@st.fragment
def render_workspace_voice(file_names_voice):
    st.markdown("<div style='font-size:14px; font-weight:600; color:#475569; margin-bottom:8px;'>Pilih & Filter Voice Aktif:</div>", unsafe_allow_html=True)
    if 'ms_voice' not in st.session_state:
        st.session_state.ms_voice = file_names_voice[:1] if file_names_voice else []

    col_f1_v, col_s1_v, col_s2_v = st.columns([3, 5, 2], vertical_alignment="bottom")
    with col_f1_v:
        pilihan_grup_v = st.selectbox("Berdasarkan Sub-Corpus:", ["Semua File"] + list(st.session_state.sub_corpora.keys()), label_visibility="collapsed", key="pg_voice") 
    with col_s2_v:
        cek_semua_v = st.checkbox("Pilih Semua", key="ca_voice")
    with col_s1_v:
        tersedia_v = file_names_voice if pilihan_grup_v == "Semua File" else [f for f in st.session_state.sub_corpora.get(pilihan_grup_v, []) if f in file_names_voice]
        if cek_semua_v:
            st.session_state.ms_voice = tersedia_v
        with st.popover(f"🎙️ Pilih File ({len([f for f in st.session_state.ms_voice if f in tersedia_v])} Terpilih)", use_container_width=True):
            temp_selection_v = []
            for f in tersedia_v:
                if st.checkbox(f, value=(f in st.session_state.ms_voice), key=f"chk_v_{f}"):
                    temp_selection_v.append(f)
            if not cek_semua_v:
                st.session_state.ms_voice = temp_selection_v
                
    selected_voice = [doc for doc in st.session_state.ms_voice if doc in tersedia_v]

    if selected_voice:
        st.markdown("---")
        tab_c_voice, tab_t_voice = st.tabs(["⚖️ Compare & Visual", "⏱️ Timeline Transkrip"])
        with tab_c_voice: 
            render_tab_compare(selected_voice, suffix="voice")
            st.markdown("<hr style='border: 2px dashed #CBD5E1; margin: 40px 0;'>", unsafe_allow_html=True)
            render_tab_visual(selected_voice, suffix="voice")   
        with tab_t_voice: 
            render_tab_transkrip(selected_voice)
    else:
        st.warning("⚠️ Silakan pilih minimal 1 file audio untuk dianalisis.")


with tab_induk_voice:
    st.info("Audio yang diunggah akan otomatis ditranskripsi menjadi teks korpus.")
    uploaded_voice = st.file_uploader(
        "Upload File Suara (Support: MP3, WAV, M4A, AAC)", 
        accept_multiple_files=True, 
        type=['mp3', 'wav', 'm4a', 'aac'],
        key=f"uploader_voice_{st.session_state.uploader_key}"
    )
    
    if uploaded_voice:
        file_voice_baru = [f for f in uploaded_voice if f.name not in st.session_state.local_files]
        if file_voice_baru:
            ada_file_baru = True
            total_voice = len(file_voice_baru)
            bar_progres_v = st.progress(0, text="Menyiapkan mesin transkripsi AI...")
            for i, file in enumerate(file_voice_baru):
                base_pct = (i / total_voice) * 100

                def hitung_pct_v(step_pct): return int(base_pct + (step_pct / total_voice))
                bar_progres_v.progress(hitung_pct_v(15), text=f"🎙️ [1/5] Mentranskripsi audio '{file.name}' (Membutuhkan waktu)...")
                raw_text, segments = transkripsi_suara_whisper(file)

                bar_progres_v.progress(hitung_pct_v(60), text=f"🧹 [2/5] Membersihkan teks hasil transkripsi '{file.name}'...")
                teks_bersih = bersihkan_teks_untuk_analisis(raw_text)    

                bar_progres_v.progress(hitung_pct_v(70), text=f"🌐 [3/5] Mengidentifikasi bahasa audio '{file.name}'...")
                try: deteksi_lang = detect(teks_bersih[:5000]) 
                except: deteksi_lang = 'en'

                if deteksi_lang not in SPACY_MODELS: deteksi_lang = 'en'               
                bar_progres_v.progress(hitung_pct_v(85), text=f"🧠 [4/5] Menyesuaikan model Linguistik '{file.name}'...")

                bar_progres_v.progress(hitung_pct_v(95), text=f"📊 [5/5] Menghitung statistik kata '{file.name}'...")
                pola_kata = r'\b[a-zA-Z0-9]+(?:-[a-zA-Z0-9]+)*\b'
                semua_kata = re.findall(pola_kata, raw_text.lower())
                
                try:
                    import nltk
                    jml_kalimat = len(nltk.sent_tokenize(raw_text[:20000]))
                except:
                    jml_kalimat = raw_text[:20000].count('.') + raw_text[:20000].count('!') + raw_text[:20000].count('?')
                
                st.session_state.local_files[file.name] = {
                    'text': raw_text, 
                    'cleaned': teks_bersih, 
                    'lang': deteksi_lang,
                    'vocab': set(semua_kata), 
                    'type': 'voice', 
                    'segments': segments, 
                    'audio_bytes': file.getvalue(),
                    'stats': {'k': jml_kalimat, 'w': len(semua_kata)}
                }
                bar_progres_v.progress(hitung_pct_v(100), text=f"✅ Transkripsi '{file.name}' berhasil!")
                st.session_state.notif_msg = f"Transkripsi audio '{file.name}' berhasil diproses!"
                st.session_state.notif_time = time.time()
                
            bar_progres_v.empty()
            st.session_state.uploader_key += 1
            st.rerun()

    voice_files = {k: v for k, v in st.session_state.local_files.items() if v.get('type') == 'voice'}
    
    if voice_files:
        file_names_voice = list(voice_files.keys())
        with st.expander("📂 Manajemen Korpus & Sub-Corpora (Voice)", expanded=False):
            t_manage_v, t_group_v = st.tabs(["📄 Daftar File", "🗂️ Kelola Sub-Corpora"])
            with t_manage_v:
                corpus_data_voice = []
                for f_name, f_data in voice_files.items():
                    grup_file = [g for g, files in st.session_state.sub_corpora.items() if f_name in files]
                    corpus_data_voice.append({
                        "Hapus": False, "Nama File": f_name, "Bahasa": f_data.get('lang', 'en').upper(),
                        "Sub-Corpus": ", ".join(grup_file) if grup_file else "Unassigned",
                        "Total Kata": f_data['stats']['w'],
                        "Kekayaan Kata (%)": round((len(f_data['vocab']) / f_data['stats']['w']) * 100, 2) if f_data['stats']['w'] > 0 else 0,
                    })
                df_corpus_voice = pd.DataFrame(corpus_data_voice)
                edited_df_voice = st.data_editor(df_corpus_voice, column_config={"Hapus": st.column_config.CheckboxColumn("❌ Hapus", default=False), "Nama File": st.column_config.TextColumn("📄 Voice", disabled=True, width="medium"), "Bahasa": st.column_config.TextColumn("🌐 BHS", disabled=True), "Sub-Corpus": st.column_config.TextColumn("🗂️ Group", disabled=True), "Kekayaan Kata (%)": st.column_config.ProgressColumn("💎 Richness", format="%.2f%%", min_value=0, max_value=100)}, hide_index=True, use_container_width=True, key="ed_voice")
                file_untuk_dihapus_v = edited_df_voice[edited_df_voice["Hapus"] == True]["Nama File"].tolist()
                if st.button("🗑️ Hapus File Terpilih", type="primary", disabled=not file_untuk_dihapus_v, key="del_voice"):
                    for fname in file_untuk_dihapus_v:
                        del st.session_state.local_files[fname]
                        for g in st.session_state.sub_corpora:
                            if fname in st.session_state.sub_corpora[g]: st.session_state.sub_corpora[g].remove(fname)
                    st.rerun()
            with t_group_v:
                c1_v, c2_v = st.columns([1, 1])
                with c1_v:
                    new_group_v = st.text_input("Buat Sub-Corpus Baru:", placeholder="Contoh: Interview_2026", key="ng_voice")
                    if st.button("➕ Tambah Grup", key="addg_voice") and new_group_v:
                        if new_group_v not in st.session_state.sub_corpora:
                            st.session_state.sub_corpora[new_group_v] = []
                            st.success(f"Grup '{new_group_v}' dibuat!")
                with c2_v:
                    target_grup_v = st.selectbox("Pilih Grup Tujuan:", list(st.session_state.sub_corpora.keys()), key="tg_voice")
                    files_to_add_v = st.multiselect("Pilih File untuk Dimasukkan:", file_names_voice, key="ms_add_voice")
                    if st.button("📥 Masukkan ke Grup", key="btn_add_voice"):
                        for f in files_to_add_v:
                            if f not in st.session_state.sub_corpora[target_grup_v]: st.session_state.sub_corpora[target_grup_v].append(f)
                        st.toast(f"Berhasil update grup {target_grup_v}!")
                        st.rerun()
        render_workspace_voice(file_names_voice)
    else:
        st.info("👋 Silakan upload file rekaman suara untuk mendapatkan transkrip dan analisisnya.")
